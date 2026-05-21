"""
lens_s3_step_report.py — S3 Strategic Pattern Intelligence Report
Project Lens | LENS-023
Model: cerebras qwen-3-235b (free, fast)
Purpose: Full quality docx of long-term patterns, historical parallels,
         structural changes, and strategic first domino analysis.
Output: YYYYMMDD_S3_Strategic_Intelligence_DC{time}.docx → Telegram
"""

import os, json, time, logging, tempfile, requests
from datetime import datetime, timezone, timedelta
from typing import Optional

logging.basicConfig(level=logging.INFO,
    format="%(asctime)s [S3-RPT] %(levelname)s %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("s3_report")

MODEL       = "gpt-oss-120b"
TEMPERATURE = 0.3
MAX_TOKENS  = 4500
TELEGRAM_CAPTION_CAP = 950


# ── Data fetch ────────────────────────────────────────────────────────────────

def fetch_s3_data() -> dict:
    from supabase import create_client
    sb = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])

    # All S3 positions — latest
    s3_all = sb.table("lens_system3_reports") \
        .select("position,report_type,summary,first_domino,patterns_found,quality_score,generated_at") \
        .order("generated_at", desc=True).limit(12).execute().data or []

    s3a = next((r for r in s3_all if r.get("position") == "S3-A"), None)
    s3b = next((r for r in s3_all if r.get("position") == "S3-B"), None)
    s3c = next((r for r in s3_all if r.get("position") == "S3-C"), None)
    s3d = next((r for r in s3_all if r.get("position") == "S3-D"), None)

    # Latest prediction
    pred = sb.table("lens_predictions") \
        .select("prediction,confidence,verification_date,created_at") \
        .order("created_at", desc=True).limit(3).execute().data or []

    # Latest S1 for context
    s1 = sb.table("lens_reports") \
        .select("domain_focus,summary,quality_score") \
        .order("generated_at", desc=True).limit(4).execute().data or []

    # Latest MA verdict
    ma = (sb.table("lens_macro_reports") \
        .select("threat_level,executive_summary,contamination_depth") \
        .order("created_at", desc=True).limit(1).execute().data or [{}])[0]

    # 30-day threat trend
    cutoff_30d = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
    trend = sb.table("lens_macro_reports") \
        .select("threat_level,created_at") \
        .gte("created_at", cutoff_30d) \
        .order("created_at", desc=True).limit(20).execute().data or []

    log.info(f"S3 report: S3-A={'yes' if s3a else 'no'}, S3-B={'yes' if s3b else 'no'}, "
             f"S3-C={'yes' if s3c else 'no'}, S3-D={'yes' if s3d else 'no'}")
    return {"s3a": s3a, "s3b": s3b, "s3c": s3c, "s3d": s3d,
            "pred": pred, "s1": s1, "ma": ma, "trend": trend}


# ── AI prompt ─────────────────────────────────────────────────────────────────

def build_s3_prompt(data: dict) -> str:
    s3a  = data["s3a"] or {}
    s3b  = data["s3b"] or {}
    s3c  = data["s3c"] or {}
    s3d  = data["s3d"] or {}
    pred = data["pred"]
    s1   = data["s1"]
    ma   = data["ma"]
    trend = data["trend"]
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # Parse patterns from S3-A
    patterns = []
    raw_pats = s3a.get("patterns_found", "[]")
    if isinstance(raw_pats, str):
        try: patterns = json.loads(raw_pats)
        except: patterns = []
    elif isinstance(raw_pats, list):
        patterns = raw_pats

    # 30-day threat distribution
    threat_counts = {}
    for t in trend:
        lv = t.get("threat_level","?")
        threat_counts[lv] = threat_counts.get(lv, 0) + 1

    prompt = f"""You are the S3 Strategic Pattern Analyst for Project Lens, an OSINT geopolitical intelligence system built for GCSP educators.

Today is {date_str}. You have received the deep analytical outputs from System 3 — the long-horizon pattern detection layer. Your task is to write a comprehensive strategic intelligence report that reads like a senior geopolitical analyst's briefing. Long analytical paragraphs, cause-effect reasoning, historical depth, no bullet points in body text. Each Part minimum 200 words.

=== CURRENT OPERATIONAL CONTEXT ===
Current Threat Level: {ma.get('threat_level', 'UNKNOWN')}
Contamination Depth: {ma.get('contamination_depth', 'UNKNOWN')}
MA Verdict: {ma.get('executive_summary', 'Not available')[:400]}

30-Day Threat Distribution: {json.dumps(threat_counts)}

=== S1 CURRENT SIGNAL (today's canary reading) ===
"""
    for r in s1:
        prompt += f"{r.get('domain_focus','?')} (quality {r.get('quality_score',0)}/10): {r.get('summary','')[:300]}\n\n"

    prompt += "\n=== S3-A — 7-DAY PATTERN INTELLIGENCE ===\n"
    if s3a:
        prompt += f"Generated: {s3a.get('generated_at','?')[:16]} | Quality: {s3a.get('quality_score',0)}\n"
        prompt += f"Summary: {s3a.get('summary','Not available')}\n\n"
        prompt += f"First Domino: {s3a.get('first_domino','Not identified')}\n\n"
        if patterns:
            prompt += f"Patterns detected ({len(patterns)}):\n"
            for p in patterns[:5]:
                if isinstance(p, dict):
                    prompt += f"  - {p.get('pattern','')}: {p.get('description','')[:200]}\n"
                elif isinstance(p, str):
                    prompt += f"  - {p[:200]}\n"
    else:
        prompt += "S3-A data not available for this cycle.\n"

    prompt += "\n=== S3-B — HISTORICAL PARALLEL ANALYSIS ===\n"
    if s3b:
        prompt += f"Generated: {s3b.get('generated_at','?')[:16]} | Quality: {s3b.get('quality_score',0)}\n"
        prompt += f"Summary: {s3b.get('summary','Not available')}\n\n"
        prompt += f"First Domino: {s3b.get('first_domino','Not identified')}\n"
    else:
        prompt += "S3-B data not available for this cycle.\n"

    prompt += "\n=== S3-D — 30-DAY STRUCTURAL CHANGE ANALYSIS ===\n"
    if s3d:
        prompt += f"Generated: {s3d.get('generated_at','?')[:16]} | Quality: {s3d.get('quality_score',0)}\n"
        prompt += f"Summary: {s3d.get('summary','Not available')}\n\n"
        prompt += f"Structural First Domino: {s3d.get('first_domino','Not identified')}\n"
    else:
        prompt += "S3-D data not available for this cycle.\n"

    prompt += "\n=== S3-C — BIAS DRIFT MONITOR (WEEKLY) ===\n"
    if s3c:
        prompt += f"Generated: {s3c.get('generated_at','?')[:16]} | Quality: {s3c.get('quality_score',0)}\n"
        prompt += f"Summary: {s3c.get('summary','Not available')}\n"
    else:
        prompt += "S3-C data not available (runs weekly on Thursdays).\n"

    if pred:
        prompt += "\n=== RECORDED PREDICTIONS ===\n"
        for p in pred[:3]:
            prompt += f"Prediction: {p.get('prediction','')[:300]}\n"
            prompt += f"Confidence: {p.get('confidence',0):.0%} | Verify by: {p.get('verification_date','?')}\n\n"

    prompt += """

=== REPORT INSTRUCTIONS ===

Write a full strategic intelligence report using exactly this structure. Each Part must be written as flowing analytical prose with multiple long paragraphs. Minimum 200 words per Part. No bullet points in body text.

PART A — THE 7-DAY PATTERN LANDSCAPE
Analyze the patterns S3-A detected over the past 7 days. What structural trends are forming? What do these patterns reveal about the trajectory of the geopolitical information environment? How do the individual patterns interconnect into a larger picture?

PART B — THE FIRST DOMINO CHAIN ANALYSIS
What is the "first domino" identified by System 3? If current patterns continue, what sequence of events becomes increasingly inevitable? Walk through the causal chain analytically. What are the trigger conditions? What would accelerate or slow this chain?

PART C — HISTORICAL PARALLEL AND PRECEDENT
What historical precedent does S3-B identify for current patterns? How closely does the current situation mirror historical cases? What does history tell us about probable outcomes? Where does the historical parallel break down and why does that matter?

PART D — STRUCTURAL CHANGE OVER 30 DAYS
What has fundamentally changed in the geopolitical information environment over the past 30 days (S3-D)? These are structural shifts — not daily noise but tectonic movements. What institutions, power relationships, or narrative architectures have been permanently altered?

PART E — ANALYTICAL DRIFT AND BIAS CHECK
What does the S3-C weekly bias monitor reveal? Are the analytical positions of the system drifting in any direction? Is the framing of certain actors becoming systematically skewed? What corrections should GCSP educators apply to their reading of this system's outputs?

PART F — STRATEGIC VERDICT AND PREDICTIONS
Synthesize all S3 findings into a strategic verdict. What is Project Lens's current assessment of the long-term trajectory? What predictions have been recorded and what evidence would confirm or deny them? What should GCSP educators watch for in the coming week?

Write in formal strategic intelligence briefing style. Use PART A, PART B format exactly for headings."""

    return prompt


# ── AI call (Cerebras) ────────────────────────────────────────────────────────

def call_cerebras(prompt: str) -> Optional[str]:
    api_key = os.environ.get("CEREBRAS_API_KEY", "")
    if not api_key:
        log.error("CEREBRAS_API_KEY not set"); return None
    for attempt in range(1, 4):
        try:
            log.info(f"S3 report calling Cerebras (attempt {attempt})")
            r = requests.post(
                "https://api.cerebras.ai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": MODEL, "messages": [{"role": "user", "content": prompt}],
                      "max_tokens": MAX_TOKENS, "temperature": TEMPERATURE},
                timeout=120)
            if r.status_code == 200:
                text = r.json()["choices"][0]["message"]["content"].strip()
                log.info(f"S3 report: {len(text)} chars generated")
                return text
            log.warning(f"Cerebras {r.status_code} attempt {attempt}: {r.text[:200]}")
            time.sleep(20 * attempt)
        except Exception as e:
            log.error(f"Cerebras call failed: {e}"); time.sleep(15)
    log.warning("Cerebras exhausted - falling back to Mistral-small")
    return call_mistral_fallback(prompt)


def call_mistral_fallback(prompt: str):
    import time as _t
    api_key = os.environ.get("MISTRAL_API_KEY", "")
    if not api_key:
        log.error("MISTRAL_API_KEY not set - no fallback available"); return None
    for attempt in range(1, 3):
        try:
            log.info(f"S3 report calling Mistral fallback (attempt {attempt})")
            r = requests.post(
                "https://api.mistral.ai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": "mistral-small-latest",
                      "messages": [{"role": "user", "content": prompt}],
                      "max_tokens": MAX_TOKENS, "temperature": TEMPERATURE},
                timeout=120)
            if r.status_code == 200:
                text = r.json()["choices"][0]["message"]["content"].strip()
                log.info(f"S3 report (Mistral fallback): {len(text)} chars generated")
                return text
            log.warning(f"Mistral fallback {r.status_code} attempt {attempt}: {r.text[:200]}")
            _t.sleep(20 * attempt)
        except Exception as e:
            log.error(f"Mistral fallback failed attempt {attempt}: {e}"); _t.sleep(15)
    return None



# ── Docx renderer ─────────────────────────────────────────────────────────────

def render_docx(report_text: str, date_str: str, s3a: dict, s3d: dict) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROJECT LENS — S3 STRATEGIC PATTERN INTELLIGENCE REPORT")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"{date_str}  |  System 3  |  Long-Horizon Analysis  |  Cerebras")
    sr.font.size = Pt(10); sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    q3a = s3a.get("quality_score", 0) if s3a else 0
    q3d = s3d.get("quality_score", 0) if s3d else 0
    dep = doc.add_paragraph()
    dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dep.add_run(f"S3-A Quality: {q3a}  |  S3-D Quality: {q3d}  |  7-day + 30-day analysis")
    dr.italic = True; dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x33, 0x66, 0x00)

    doc.add_paragraph()

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph(); continue
        if line.startswith("PART ") and "—" in line:
            doc.add_heading(line, level=1)
        elif line.isupper() and len(line) < 80 and ":" not in line:
            doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    time_str = datetime.now(timezone.utc).strftime("%H%M")
    fname = f"{date_str.replace('-','')}_S3_Strategic_Intelligence_DC{time_str}.docx"
    tmp_path = os.path.join(tempfile.gettempdir(), fname)
    doc.save(tmp_path)
    log.info(f"S3 docx saved: {tmp_path}")
    return tmp_path


# ── Telegram ──────────────────────────────────────────────────────────────────

def send_telegram_doc(docx_path: str, caption: str) -> bool:
    token   = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        log.warning("Telegram keys not set"); return False
    try:
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        with open(docx_path, "rb") as f:
            r = requests.post(url, data={"chat_id": chat_id, "caption": caption,
                "parse_mode": "HTML"}, files={"document": (
                    os.path.basename(docx_path), f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )}, timeout=60)
        if r.status_code == 200:
            log.info("S3 docx sent to Telegram"); return True
        log.error(f"Telegram error: {r.status_code}"); return False
    except Exception as e:
        log.error(f"Telegram send failed: {e}"); return False


def send_telegram_text(text: str) -> bool:
    token   = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id: return False
    try:
        r = requests.post(f"https://api.telegram.org/bot{token}/sendMessage",
            json={"chat_id": chat_id, "text": text, "parse_mode": "HTML"}, timeout=10)
        return r.status_code == 200
    except: return False


# ── Entry point ───────────────────────────────────────────────────────────────

def run_s3_report() -> dict:
    start = time.time()
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    log.info(f"=== S3 STRATEGIC REPORT START | {date_str} ===")

    try:
        data = fetch_s3_data()
    except Exception as e:
        log.error(f"Data fetch failed: {e}")
        return {"status": "ERROR", "error": str(e)}

    if not data["s3a"] and not data["s3d"]:
        log.warning("No S3 reports found")
        return {"status": "NO_DATA"}

    prompt = build_s3_prompt(data)
    report_text = call_cerebras(prompt)
    if not report_text:
        return {"status": "AI_FAILED"}

    try:
        docx_path = render_docx(report_text, date_str, data["s3a"] or {}, data["s3d"] or {})
    except Exception as e:
        log.error(f"Docx render failed: {e}")
        return {"status": "DOCX_FAILED", "error": str(e)}

    s3a = data["s3a"] or {}
    dom = (s3a.get("first_domino","") or "")[:120]
    pats_raw = s3a.get("patterns_found","[]")
    try:
        pats = json.loads(pats_raw) if isinstance(pats_raw, str) else pats_raw
        pat_count = len(pats) if isinstance(pats, list) else 0
    except: pat_count = 0

    intro = (
        f"📚 <b>S3 Strategic Pattern Report — {date_str}</b>\n"
        f"{pat_count} patterns detected | 7-day + 30-day horizon\n"
        f"<i>Full strategic analysis attached — patterns, historical parallels, "
        f"first domino chain, and predictions</i>"
    )
    send_telegram_text(intro)
    time.sleep(1)

    caption = (
        f"📚 S3 Strategic Pattern Intelligence Report — {date_str}\n"
        f"System 3 | {pat_count} patterns | Long-horizon analysis\n"
        f"Parts: 7-day Patterns | First Domino | Historical Parallel | "
        f"Structural Change | Drift Check | Strategic Verdict"
    )[:TELEGRAM_CAPTION_CAP]

    sent = send_telegram_doc(docx_path, caption)
    elapsed = round(time.time() - start, 1)
    log.info(f"=== S3 REPORT COMPLETE | sent={sent} | {elapsed}s ===")
    return {"status": "COMPLETE" if sent else "SEND_FAILED", "elapsed": elapsed}


if __name__ == "__main__":
    from dotenv import load_dotenv; load_dotenv()
    result = run_s3_report()
    print(result)
