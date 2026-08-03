"""
lens_compendium.py
Project Lens — Intelligence Compendium (Free Tier)
Model: Groq llama-3.3-70b-versatile (light synthesis + formatting)
Schedule: 1x daily at 02:30 UTC (10:30 PM DC EDT)
Cost: $0/run

6 Sections:
  1. HOW THE INFORMATION IS BEING SHAPED (S2 injection detail)
  2. ENTITY INTELLIGENCE (named actors, legitimacy, frequency)
  3. S2-F OPERATION ALERTS (Watch/Clarity/Verification findings)
  4. SYSTEM 3 PATTERN BRIEF (all S3 positions)
  5. 7-DAY TREND SUMMARY (daily intelligence trends)
  6. PREDICTION TRACKER (active S3-A predictions)

Guard: preflight + retry (3x) + failure alert
Output: docx → Telegram
"""

import logging
import os
import re
import sys
import tempfile
import time
from datetime import datetime, timedelta, timezone

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [COMPENDIUM] %(levelname)s %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("COMPENDIUM")

# ── Constants ─────────────────────────────────────────────────────────────────
LOOKBACK_HOURS = 24
MAX_TOKENS     = 3000
TEMPERATURE    = 0.2
MAX_RETRIES    = 3
CAPTION_CAP    = 950

# Intro synthesis system prompt -- module scope so the probe fixture can
# import it instead of keeping a second copy (fixtures import prompts,
# never duplicate them).
INTRO_SYSTEM_PROMPT = (
    "You are Project Lens. Write a 3-sentence executive introduction "
    "for today's Intelligence Compendium. Concise, intelligence-briefing "
    "tone. No predictions. PHI-003: use Office names (Xi Office, Trump Office etc)."
)


# ── Guard: Preflight ──────────────────────────────────────────────────────────
def preflight_check() -> bool:
    required = ["SUPABASE_URL", "SUPABASE_SERVICE_KEY",
                "TELEGRAM_BOT_TOKEN", "TELEGRAM_CHAT_ID", "GROQ_API_KEY"]
    missing = [k for k in required if not os.environ.get(k)]
    if missing:
        log.error(f"PREFLIGHT FAILED — missing: {missing}")
        return False
    log.info("PREFLIGHT OK")
    return True


# ── Clients ───────────────────────────────────────────────────────────────────
def get_supabase():
    from supabase import create_client
    return create_client(
        os.environ["SUPABASE_URL"],
        os.environ["SUPABASE_SERVICE_KEY"]
    )


def get_groq():
    from groq import Groq
    key = os.environ.get("GROQ_S2DGCOM_API_KEY") or os.environ["GROQ_API_KEY"]
    return Groq(api_key=key)


# ── Section 1: HOW THE INFORMATION IS BEING SHAPED ───────────────────────────
def fetch_s2_injection_detail(sb) -> list:
    """All injection reports last 24h."""
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=LOOKBACK_HOURS)).isoformat()
    try:
        r = sb.table("injection_reports") \
            .select("analyst,injection_type,evidence,confidence_score,flagged_phrases,cycle,created_at") \
            .gte("created_at", cutoff) \
            .order("confidence_score", desc=True) \
            .limit(30) \
            .execute()
        log.info(f"S2 injections: {len(r.data or [])} findings")
        return r.data or []
    except Exception as e:
        log.warning(f"S2 injection fetch failed: {e}")
        return []


def build_section1(injections: list) -> str:
    """Section 1: How the information is being shaped."""
    lines = ["=" * 60]
    lines.append("SECTION 1 — HOW THE INFORMATION IS BEING SHAPED")
    lines.append("=" * 60)
    lines.append(f"Total S2 findings: {len(injections)} | Window: last 24h")
    lines.append("")

    if not injections:
        lines.append("No S2 injection findings in this window.")
        return "\n".join(lines)

    # Group by injection type
    by_type = {}
    for inj in injections:
        itype = inj.get("injection_type", "UNKNOWN")
        if itype not in by_type:
            by_type[itype] = []
        by_type[itype].append(inj)

    # Top injection type
    top_type = max(by_type, key=lambda k: len(by_type[k]))
    top_conf = max((i.get("confidence_score", 0) for i in injections), default=0)
    lines.append(f"Dominant injection method: {top_type} ({len(by_type[top_type])} findings, max conf {top_conf:.2f})")
    lines.append("")

    for itype, items in sorted(by_type.items(), key=lambda x: -len(x[1])):
        lines.append(f"▶ {itype} ({len(items)} findings):")
        for item in items[:3]:  # top 3 per type
            evidence_raw = item.get("evidence") or {}
            evidence = str(evidence_raw.get("key_claims", evidence_raw))[:300] if isinstance(evidence_raw, dict) else str(evidence_raw)[:300]
            flagged = item.get("flagged_phrases", [])
            conf = item.get("confidence_score", 0)
            analyst = item.get("analyst", "")
            lines.append(f"  [{analyst}] conf={conf:.2f}")
            if evidence:
                lines.append(f"  Evidence: {evidence}")
            if flagged and isinstance(flagged, list) and flagged:
                lines.append(f"  Trigger language: {', '.join(str(f) for f in flagged[:3])}")
            lines.append("")

    # GAP findings separately
    gap_items = [i for i in injections if i.get("injection_type") == "GAP_ANALYSIS"]
    if gap_items:
        lines.append("── WHAT THE CANARY MISSED (Broken Window) ──")
        for g in gap_items[:3]:
            evidence_raw = g.get("evidence") or {}
            if isinstance(evidence_raw, dict):
                claims = evidence_raw.get("key_claims", [])
                evidence = str(claims[0] if claims else evidence_raw)[:400]
            else:
                evidence = str(evidence_raw)[:400]
            lines.append(f"{evidence}")
            lines.append("")

    return "\n".join(lines)


# ── Section 2: ENTITY INTELLIGENCE ───────────────────────────────────────────
def fetch_entities(sb) -> list:
    """Most active entities by mention count."""
    try:
        r = sb.table("lens_entities") \
            .select("entity_type,name,canonical_name,primary_outlet,affiliations,"
                    "total_mentions,last_seen") \
            .order("total_mentions", desc=True) \
            .limit(20) \
            .execute()
        log.info(f"Entities: {len(r.data or [])} records")
        return r.data or []
    except Exception as e:
        log.warning(f"Entities fetch failed: {e}")
        return []


def build_section2(entities: list) -> str:
    """Section 2: Entity intelligence."""
    lines = ["=" * 60]
    lines.append("SECTION 2 — ENTITY INTELLIGENCE")
    lines.append("=" * 60)

    if not entities:
        lines.append("No entity data available yet.")
        return "\n".join(lines)

    lines.append(f"Total tracked entities: {len(entities)}")
    lines.append("")

    # Group by type
    by_type = {}
    for e in entities:
        etype = e.get("entity_type", "UNKNOWN")
        if etype not in by_type:
            by_type[etype] = []
        by_type[etype].append(e)

    for etype, items in sorted(by_type.items()):
        lines.append(f"▶ {etype}:")
        for item in items[:5]:
            name = item.get("canonical_name") or item.get("name", "")
            mentions = item.get("total_mentions", 0)
            outlet = (item.get("primary_outlet") or "")
            last_seen = (item.get("last_seen") or "")[:10]
            affiliations = item.get("affiliations", [])
            affil_str = ""
            if affiliations and isinstance(affiliations, list):
                affil_str = f" | Affiliations: {', '.join(str(a) for a in affiliations[:2])}"
            lines.append(f"  {name}: {mentions} mentions | Last: {last_seen}{affil_str}")
            if outlet:
                lines.append(f"    Primary outlet: {outlet}")
        lines.append("")

    return "\n".join(lines)


# ── Section 3: S2-F OPERATION ALERTS ─────────────────────────────────────────
def fetch_s2f_data(sb) -> dict:
    """S2-F detections and drift findings."""
    result = {"detections": [], "drift": []}
    cutoff_24h = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
    cutoff_45d = (datetime.now(timezone.utc) - timedelta(days=45)).isoformat()

    try:
        r = sb.table("lens_operation_detections") \
            .select("state_actor_lens,operations_detected,operation_count,"
                    "confidence,food_for_thought,scored_at") \
            .gte("scored_at", cutoff_24h) \
            .eq("not_applicable", False) \
            .order("confidence", desc=True) \
            .limit(20) \
            .execute()
        result["detections"] = r.data or []
        log.info(f"S2-F detections: {len(result['detections'])} in last 24h")
    except Exception as e:
        log.warning(f"S2-F detections fetch failed: {e}")

    try:
        r = sb.table("lens_drift_findings") \
            .select("state_actor_lens,finding_confidence,finding_phrasing,sample_size,created_at") \
            .in_("finding_confidence", ["HIGH", "MEDIUM"]) \
            .gte("created_at", cutoff_45d) \
            .order("finding_confidence", desc=True) \
            .limit(10) \
            .execute()
        result["drift"] = r.data or []
        log.info(f"S2-F drift: {len(result['drift'])} HIGH/MEDIUM findings")
    except Exception as e:
        log.warning(f"S2-F drift fetch failed: {e}")

    return result


def build_section3(s2f: dict) -> str:
    """Section 3: S2-F operation alerts."""
    import json as _json
    lines = ["=" * 60]
    lines.append("SECTION 3 — S2-F PRETENSE OPERATION ALERTS")
    lines.append("=" * 60)

    detections = s2f.get("detections", [])
    drift = s2f.get("drift", [])

    if not detections and not drift:
        lines.append("No S2-F operation findings yet.")
        lines.append("(S2-F pipeline is new — findings accumulate over 7-45 days)")
        return "\n".join(lines)

    if drift:
        lines.append(f"CONFIRMED PATTERNS ({len(drift)} HIGH/MEDIUM, last 45 days):")
        lines.append("")
        for f in drift:
            lens = f.get("state_actor_lens", "")
            conf = f.get("finding_confidence", "")
            phrasing = (f.get("finding_phrasing") or "")[:400]
            sample = f.get("sample_size", 0)
            created = (f.get("created_at") or "")[:10]
            lines.append(f"  [{conf}] {lens} | {sample} articles | {created}")
            lines.append(f"  {phrasing}")
            lines.append("")

    if detections:
        lines.append(f"RECENT DETECTIONS ({len(detections)} article detections, last 24h):")
        lines.append("")
        for det in detections[:10]:
            lens = det.get("state_actor_lens", "")
            conf = det.get("confidence", 0)
            count = det.get("operation_count", 0)
            fft = (det.get("food_for_thought") or "")[:200]
            ops = det.get("operations_detected") or []
            if isinstance(ops, str):
                try:
                    ops = _json.loads(ops)
                except Exception:
                    ops = []
            lines.append(f"  {lens} | conf={conf:.2f} | {count} operations detected")
            for op in ops[:3]:
                op_id = op.get("id", "")
                op_name = op.get("name", "")
                evidence = (op.get("evidence_phrase") or "")[:150]
                lines.append(f"    {op_id}: {op_name}")
                lines.append(f"      Evidence: {evidence}")
            if fft:
                lines.append(f"  ❓ {fft}")
            lines.append("")

    return "\n".join(lines)


# ── Section 4: SYSTEM 3 PATTERN BRIEF ────────────────────────────────────────
def fetch_s3_all(sb) -> dict:
    """Latest report per S3 position."""
    positions = ["S3-A", "S3-B", "S3-C", "S3-D", "S3-E"]
    results = {}
    for pos in positions:
        try:
            r = sb.table("lens_system3_reports") \
                .select("position,report_type,summary,first_domino,patterns_found,"
                        "structural_trends,signals_to_watch,quality_score,generated_at") \
                .eq("position", pos) \
                .order("generated_at", desc=True) \
                .limit(1) \
                .execute()
            if r.data:
                results[pos] = r.data[0]
        except Exception as e:
            log.warning(f"S3 {pos} fetch failed: {e}")
    log.info(f"S3: {list(results.keys())} positions with data")
    return results


def build_section4(s3: dict) -> str:
    """Section 4: S3 pattern brief."""
    lines = ["=" * 60]
    lines.append("SECTION 4 — SYSTEM 3 PATTERN BRIEF")
    lines.append("=" * 60)

    if not s3:
        lines.append("No S3 pattern data available.")
        return "\n".join(lines)

    pos_labels = {
        "S3-A": "7-Day Pattern Detector (Groq)",
        "S3-B": "True History Analyst (Gemini)",
        "S3-C": "Bias Drift Monitor (Cohere)",
        "S3-D": "Long-term Structural (Cerebras)",
        "S3-E": "Self-Check (Ollama LOCAL)",
    }

    for pos in ["S3-A", "S3-B", "S3-C", "S3-D", "S3-E"]:
        label = pos_labels.get(pos, pos)
        if pos not in s3:
            lines.append(f"▶ {pos} — {label}: No data (check cadence)")
            lines.append("")
            continue

        rep = s3[pos]
        summary = (rep.get("summary") or "")[:500]
        first_domino = (rep.get("first_domino") or "")[:250]
        signals = (rep.get("signals_to_watch") or "")[:250]
        structural = (rep.get("structural_trends") or "")[:300]
        quality = rep.get("quality_score", 0)
        generated = (rep.get("generated_at") or "")[:16]

        lines.append(f"▶ {pos} — {label}")
        lines.append(f"  Quality: {quality} | Generated: {generated}")
        lines.append(f"  Summary: {summary}")
        if first_domino:
            lines.append(f"  First domino: {first_domino}")
        if structural:
            lines.append(f"  Structural trend: {structural}")
        if signals:
            lines.append(f"  Signals to watch: {signals}")
        lines.append("")

    return "\n".join(lines)


# ── Section 5: 7-DAY TREND SUMMARY ───────────────────────────────────────────
def fetch_7day_trend(sb) -> dict:
    """7 days of macro + lens reports for trend analysis."""
    cutoff = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
    result = {"macro": [], "lens": []}
    try:
        r = sb.table("lens_macro_reports") \
            .select("threat_level,quality_score,cui_bono_synthesis,created_at") \
            .gte("created_at", cutoff) \
            .order("created_at", desc=True) \
            .limit(14) \
            .execute()
        result["macro"] = r.data or []
    except Exception as e:
        log.warning(f"Macro trend fetch failed: {e}")

    try:
        r = sb.table("injection_reports") \
            .select("injection_type,confidence_score,created_at") \
            .gte("created_at", cutoff) \
            .order("created_at", desc=True) \
            .limit(100) \
            .execute()
        result["lens"] = r.data or []
    except Exception as e:
        log.warning(f"Lens trend fetch failed: {e}")

    log.info(f"7-day trend: {len(result['macro'])} macro + {len(result['lens'])} injection records")
    return result


def build_section5(trend: dict) -> str:
    """Section 5: 7-day trend summary."""
    lines = ["=" * 60]
    lines.append("SECTION 5 — 7-DAY TREND SUMMARY")
    lines.append("=" * 60)

    macro = trend.get("macro", [])
    injections = trend.get("lens", [])

    if not macro:
        lines.append("No trend data available yet (need 7 days of data).")
        return "\n".join(lines)

    # Threat level history
    lines.append("Threat level history (newest first):")
    threat_counts = {}
    for m in macro:
        level = m.get("threat_level", "UNKNOWN")
        date = (m.get("created_at") or "")[:10]
        quality = m.get("quality_score", 0)
        lines.append(f"  {date}: {level} (quality {quality:.2f})")
        threat_counts[level] = threat_counts.get(level, 0) + 1
    lines.append("")

    # Dominant threat
    if threat_counts:
        dominant = max(threat_counts, key=threat_counts.get)
        lines.append(f"Dominant threat level this week: {dominant} ({threat_counts[dominant]}/{len(macro)} cycles)")
    lines.append("")

    # Injection type frequency
    if injections:
        type_freq = {}
        for inj in injections:
            itype = inj.get("injection_type", "UNKNOWN")
            type_freq[itype] = type_freq.get(itype, 0) + 1

        lines.append("Injection type frequency (7 days):")
        for itype, count in sorted(type_freq.items(), key=lambda x: -x[1]):
            lines.append(f"  {itype}: {count} findings")
        lines.append("")

        # Top injection this week
        top = max(type_freq, key=type_freq.get)
        lines.append(f"Most persistent injection this week: {top}")

    return "\n".join(lines)


# ── Section 6: PREDICTION TRACKER ────────────────────────────────────────────
def fetch_predictions(sb) -> list:
    """Active predictions from S3-A."""
    try:
        r = sb.table("lens_predictions") \
            .select("*") \
            .order("created_at", desc=True) \
            .limit(20) \
            .execute()
        log.info(f"Predictions: {len(r.data or [])} records")
        return r.data or []
    except Exception as e:
        log.warning(f"Predictions fetch failed: {e}")
        return []


def build_section6(predictions: list) -> str:
    """Section 6: Prediction tracker."""
    lines = ["=" * 60]
    lines.append("SECTION 6 — PREDICTION TRACKER")
    lines.append("=" * 60)

    if not predictions:
        lines.append("No active predictions yet.")
        lines.append("(S3-A starts generating predictions after sufficient data accumulates)")
        lines.append("(S4-B Outcome Verifier: planned July 2026)")
        return "\n".join(lines)

    now = datetime.now(timezone.utc)
    lines.append(f"Active predictions: {len(predictions)}")
    lines.append("")

    for pred in predictions:
        pred_text = (pred.get("prediction") or str(pred))[:300]
        created = (pred.get("created_at") or "")[:10]
        status = pred.get("status", "pending")
        verify_by = pred.get("verify_by") or ""

        # Calculate days remaining
        days_str = ""
        if verify_by:
            try:
                vdate = datetime.fromisoformat(verify_by.replace("Z", "+00:00"))
                days_left = (vdate - now).days
                days_str = f" | {days_left} days to verification"
            except Exception:
                pass

        lines.append(f"  [{status.upper()}] {created}{days_str}")
        lines.append(f"  {pred_text}")
        lines.append("")

    return "\n".join(lines)


# ── Groq synthesis (light) ────────────────────────────────────────────────────
def synthesize_intro(sections_text: str) -> str:
    """Ask Groq for a 3-sentence executive intro for the compendium."""
    try:
        client = get_groq()
        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": INTRO_SYSTEM_PROMPT},
                {"role": "user", "content": (
                    f"Based on this data summary, write the 3-sentence intro:\n\n"
                    f"{sections_text[:3000]}"
                )},
            ],
            max_tokens=200,
            temperature=0.2,
            timeout=60,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        log.warning(f"Groq intro synthesis failed: {e}")
        return "Project Lens Intelligence Compendium — full daily intelligence package."


# ── DOCX renderer ─────────────────────────────────────────────────────────────
def render_docx(sections: list, intro: str, date_str: str) -> str:
    """Render all sections to docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECT LENS — INTELLIGENCE COMPENDIUM")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run(f"{date_str}  |  6-Section Daily Package  |  Free Tier")
    sub.font.size = Pt(10)
    sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Executive intro
    intro_p = doc.add_paragraph()
    intro_run = intro_p.add_run(intro)
    intro_run.font.size = Pt(11)
    intro_run.italic = True

    doc.add_paragraph()

    # All sections
    for section_text in sections:
        lines = section_text.split("\n")
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                doc.add_paragraph()
                continue
            if line_stripped.startswith("=") and len(set(line_stripped)) == 1:
                continue  # skip separator lines
            if line_stripped.startswith("SECTION ") and "—" in line_stripped:
                p = doc.add_heading(line_stripped, level=1)
            elif line_stripped.startswith("▶ "):
                p = doc.add_heading(line_stripped[2:], level=2)
            elif line_stripped.startswith("──"):
                p = doc.add_heading(line_stripped.strip("─ "), level=3)
            else:
                p = doc.add_paragraph(line_stripped)
                p.paragraph_format.space_after = Pt(4)

        # Page break between sections
        doc.add_page_break()

    fname = f"{date_str.replace('-', '')}_ProjectLens_Compendium_DC2230.docx"
    tmp_path = os.path.join(tempfile.gettempdir(), fname)
    doc.save(tmp_path)
    log.info(f"DOCX rendered: {tmp_path}")
    return tmp_path


# ── Telegram ──────────────────────────────────────────────────────────────────
def send_telegram(docx_path: str, caption: str) -> bool:
    import requests
    token = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    url = f"https://api.telegram.org/bot{token}/sendDocument"
    try:
        with open(docx_path, "rb") as f:
            resp = requests.post(
                url,
                data={"chat_id": chat_id, "caption": caption},
                files={"document": (os.path.basename(docx_path), f,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=60,
            )
        if resp.status_code == 200:
            log.info("Telegram: OK")
            return True
        log.error(f"Telegram failed: {resp.status_code}")
        return False
    except Exception as e:
        log.error(f"Telegram error: {e}")
        return False


def send_alert(message: str) -> None:
    import requests
    token = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        return
    try:
        requests.post(
            f"https://api.telegram.org/bot{token}/sendMessage",
            data={"chat_id": chat_id, "text": f"⚠️ COMPENDIUM FAILED\n{message}"},
            timeout=30,
        )
    except Exception:
        pass


# ── Main ──────────────────────────────────────────────────────────────────────
def run_compendium(dry_run: bool = False) -> dict:
    start = time.time()
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    log.info(f"=== COMPENDIUM START | {date_str.replace('-','')} | dry_run={dry_run} ===")

    if not preflight_check():
        send_alert("Preflight failed")
        return {"status": "FAILED", "reason": "preflight"}

    try:
        sb = get_supabase()
        injections  = fetch_s2_injection_detail(sb)
        entities    = fetch_entities(sb)
        s2f         = fetch_s2f_data(sb)
        s3          = fetch_s3_all(sb)
        trend       = fetch_7day_trend(sb)
        predictions = fetch_predictions(sb)
    except Exception as e:
        msg = f"DB fetch failed: {e}"
        log.error(msg)
        send_alert(msg)
        return {"status": "FAILED", "reason": msg}

    # Build all 6 sections
    sections = [
        build_section1(injections),
        build_section2(entities),
        build_section3(s2f),
        build_section4(s3),
        build_section5(trend),
        build_section6(predictions),
    ]

    # Light Groq synthesis for intro
    if not dry_run:
        intro = synthesize_intro("\n\n".join(sections[:2]))
    else:
        intro = "Dry-run mode — intro synthesis skipped."

    if dry_run:
        total_chars = sum(len(s) for s in sections)
        log.info(f"DRY RUN — total sections chars: {total_chars}")
        return {"status": "DRY_RUN", "sections_chars": total_chars}

    # Render docx
    try:
        docx_path = render_docx(sections, intro, date_str)
    except Exception as e:
        msg = f"DOCX render failed: {e}"
        log.error(msg)
        send_alert(msg)
        return {"status": "FAILED", "reason": msg}

    # Build caption
    caption = (
        f"📦 Project Lens Intelligence Compendium — {date_str}\n"
        f"6 sections: HOW SHAPED | Entities | S2-F Alerts | S3 Patterns | 7-Day Trend | Predictions\n"
        f"S2 findings: {len(injections)} | Entities: {len(entities)} | "
        f"S3 positions: {len(s3)}"
    )[:CAPTION_CAP]

    sent = send_telegram(docx_path, caption)
    elapsed = round(time.time() - start, 1)
    log.info(f"=== COMPENDIUM DONE | sent={sent} | {elapsed}s ===")

    return {
        "status": "OK",
        "elapsed_s": elapsed,
        "sections": 6,
        "telegram_sent": sent,
    }


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    from dotenv import load_dotenv
    load_dotenv()

    result = run_compendium(dry_run=args.dry_run)
    import json
    print(json.dumps(result, indent=2))
    sys.exit(0 if result["status"] in ("OK", "DRY_RUN") else 1)
