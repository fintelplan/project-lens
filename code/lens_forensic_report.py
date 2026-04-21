"""
lens_forensic_report.py
Project Lens — Opus 4.7 Forensic Research Report

Built per LENS-016 4-decision spec + LENS-017 operator decisions:
  - Model: claude-opus-4-7 (adaptive)
  - Frequency: 1x/day via cron '0 2 * * *' UTC (= 09:00 Thai)
  - Length: 4-5 pages, ~2500 words
  - Detection depth: exhaustive (every pattern, not top-N)
  - Structure: Detection / Recovery / Food for Thought / References
  - Tone: HYBRID — Parts 1-2 intelligence-briefing formal,
                   Part 3 investigative-journalism narrative,
                   NO predictions anywhere
  - Output: DOCX via python-docx (LR-067, NOT Node.js)
  - Filename: YYYY-MM-DD_ProjectLens_Forensic_DC2200.docx
  - Telegram: file + caption <=950 chars
  - Path B citations: full reference pool to LLM, post-processing validator strips bad REF IDs

Cost target: ~$7-9/month (verified vs Opus 4.7 pricing: $5/MTok in, $25/MTok out)

Usage:
    python code/lens_forensic_report.py              # full run (real API call, ~$0.24)
    python code/lens_forensic_report.py --dry-run    # fetch + prompt build, NO API call
"""

import argparse
import json
import logging
import os
import re
import sys
import tempfile
import time
from datetime import datetime, timedelta, timezone
from typing import Any

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [FORENSIC] %(levelname)s %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("FORENSIC")

# ── Constants ─────────────────────────────────────────────────────────────────
MODEL = "claude-opus-4-7"
MAX_TOKENS = 4096          # ~2500 words target; cap = safety
LOOKBACK_HOURS = 24        # forensic window
TELEGRAM_CAPTION_CAP = 950 # per LENS-016 audit (74 char buffer under 1024)


# ── Clients ───────────────────────────────────────────────────────────────────

def get_supabase():
    from supabase import create_client
    url = os.environ.get("SUPABASE_URL", "")
    key = os.environ.get("SUPABASE_SERVICE_KEY", "")
    if not url or not key:
        raise RuntimeError("SUPABASE credentials missing")
    return create_client(url, key)


def get_anthropic():
    import anthropic
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY missing")
    return anthropic.Anthropic(api_key=api_key)


# ── Evidence fetchers (LR-076 schema-confirmed via LENS-017 SQL audit) ────────

def fetch_macro_reports(sb, hours: int = LOOKBACK_HOURS) -> list:
    """Most recent MA reports in the lookback window. Usually 2 (morning + afternoon)."""
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=hours)).isoformat()
    try:
        r = (
            sb.table("lens_macro_reports")
              .select("id,run_id,cycle,threat_level,executive_summary,key_findings,"
                      "actors_named,actors_of_concern,adversary_narrative_summary,"
                      "manufactured_narratives,gcsp_implications,intelligence_gaps,"
                      "cui_bono_synthesis,contamination_depth,quality_score,created_at")
              .gte("created_at", cutoff)
              .order("created_at", desc=True)
              .limit(4)
              .execute()
        )
        rows = r.data or []
        log.info(f"MA: {len(rows)} reports in last {hours}h")
        return rows
    except Exception as e:
        log.warning(f"MA fetch failed: {e}")
        return []


def fetch_injection_reports(sb, hours: int = LOOKBACK_HOURS) -> list:
    """All S2 injection findings in the lookback window."""
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=hours)).isoformat()
    try:
        r = (
            sb.table("injection_reports")
              .select("analyst,injection_type,evidence,confidence_score,"
                      "flagged_phrases,cycle,run_id,created_at")
              .gte("created_at", cutoff)
              .order("confidence_score", desc=True)
              .limit(80)
              .execute()
        )
        rows = r.data or []
        log.info(f"Injections: {len(rows)} S2 findings in last {hours}h")
        return rows
    except Exception as e:
        log.warning(f"Injection fetch failed: {e}")
        return []


def fetch_lens_reports(sb, hours: int = LOOKBACK_HOURS) -> list:
    """S1 lens reports in the lookback window."""
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=hours)).isoformat()
    try:
        r = (
            sb.table("lens_reports")
              .select("domain_focus,summary,food_for_thought,quality_score,"
                      "cycle,generated_at")
              .gte("generated_at", cutoff)
              .order("generated_at", desc=True)
              .limit(20)
              .execute()
        )
        rows = r.data or []
        log.info(f"S1: {len(rows)} lens reports in last {hours}h")
        return rows
    except Exception as e:
        log.warning(f"S1 fetch failed: {e}")
        return []


def fetch_system3_latest(sb) -> dict:
    """Most recent S3 report per position (no lookback — S3 may be stale; we use what exists).

    Returns {position: row} for whatever positions have data, e.g. {S3-A: {...}, S3-D: {...}}.
    Empty dict if S3 has no data at all.
    """
    out = {}
    for pos in ("S3-A", "S3-B", "S3-C", "S3-D", "S3-E"):
        try:
            r = (
                sb.table("lens_system3_reports")
                  .select("position,report_type,summary,first_domino,patterns_found,"
                          "structural_trends,signals_to_watch,quality_score,generated_at")
                  .eq("position", pos)
                  .order("generated_at", desc=True)
                  .limit(1)
                  .execute()
            )
            if r.data:
                out[pos] = r.data[0]
        except Exception as e:
            log.warning(f"S3 fetch failed for {pos}: {e}")
    log.info(f"S3: positions with data = {sorted(out.keys()) or 'NONE'}")
    return out


def fetch_reference_pool(sb) -> list:
    """All article references collected today.

    For Path B citations: pass ref_id + title + source_name + domain to the LLM
    so it can cite REF IDs by semantic relevance. The validator (later) strips
    any cited REF IDs that aren't in this pool.
    """
    today = datetime.now(timezone.utc).date().isoformat()
    try:
        r = (
            sb.table("lens_article_refs")
              .select("ref_id,title,source_name,domain")
              .eq("collected_date", today)
              .order("ref_id", desc=False)
              .execute()
        )
        rows = r.data or []
        log.info(f"References: {len(rows)} articles collected today ({today})")
        return rows
    except Exception as e:
        log.warning(f"Reference pool fetch failed: {e}")
        return []


# ── Helpers ───────────────────────────────────────────────────────────────────

def _safe_jsonb(value) -> Any:
    """Coerce a Supabase JSONB column value into a Python dict/list/etc.

    Postgrest usually returns these already-parsed, but defensive against
    string-encoded edge cases.
    """
    if value is None:
        return None
    if isinstance(value, (dict, list)):
        return value
    if isinstance(value, str):
        try:
            return json.loads(value)
        except Exception:
            return value
    return value


def _truncate(s: str, n: int) -> str:
    if not s:
        return ""
    s = str(s)
    if len(s) <= n:
        return s
    return s[: n - 3].rstrip() + "..."


# ── Prompt builder ────────────────────────────────────────────────────────────

def build_prompt(macros: list, injections: list, lens_reports: list,
                 s3_latest: dict, references: list) -> str:
    """Build the forensic report prompt per LENS-016 spec.

    Structure: Part 1 Detection (formal) / Part 2 Recovery (formal) /
               Part 3 Food for Thought (narrative, no predictions) /
               Part 4 References (auto-rendered post-LLM, but mentioned in prompt
                                  so LLM cites inline)
    """
    today_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # ── MA evidence (most recent first) ──
    ma_text = ""
    if macros:
        ma_text = "MISSION ANALYST REPORTS (most recent first):\n"
        for m in macros:
            cui = _safe_jsonb(m.get("cui_bono_synthesis")) or {}
            kfs = _safe_jsonb(m.get("key_findings")) or []
            actors = _safe_jsonb(m.get("actors_of_concern")) or []
            mn = _safe_jsonb(m.get("manufactured_narratives")) or []
            gaps = m.get("intelligence_gaps") or "(none recorded)"

            kf_lines = []
            if isinstance(kfs, list):
                for kf in kfs[:5]:
                    if isinstance(kf, dict):
                        finding = kf.get("finding") or kf.get("title") or ""
                        sig = kf.get("significance") or ""
                        kf_lines.append(f"  - {_truncate(finding, 200)}"
                                        + (f" [{_truncate(sig, 120)}]" if sig else ""))
                    else:
                        kf_lines.append(f"  - {_truncate(str(kf), 200)}")

            mn_lines = []
            if isinstance(mn, list):
                for n in mn[:4]:
                    if isinstance(n, dict):
                        mn_lines.append(f"  - {_truncate(n.get('narrative') or n.get('claim') or str(n), 180)}")
                    else:
                        mn_lines.append(f"  - {_truncate(str(n), 180)}")

            actor_lines = []
            if isinstance(actors, list):
                for a in actors[:6]:
                    if isinstance(a, dict):
                        actor_lines.append(f"  - {a.get('actor') or a.get('name') or str(a)}: "
                                           f"{_truncate(a.get('role') or a.get('rationale') or '', 100)}")
                    else:
                        actor_lines.append(f"  - {_truncate(str(a), 120)}")

            cui_text = ""
            if isinstance(cui, dict):
                cui_text = (
                    f"  Primary beneficiary: {cui.get('primary_beneficiary', '?')}\n"
                    f"  Convergence: {cui.get('convergence', '?')}\n"
                    f"  Evidence: {_truncate(str(cui.get('evidence', '')), 250)}\n"
                )

            ma_text += (
                f"\n--- MA report run_id={m.get('run_id')} "
                f"cycle={m.get('cycle')} threat={m.get('threat_level')} "
                f"depth={m.get('contamination_depth')} quality={m.get('quality_score')} ---\n"
                f"Executive summary: {_truncate(m.get('executive_summary'), 800)}\n"
                f"Adversary narrative: {_truncate(m.get('adversary_narrative_summary'), 500)}\n"
            )
            if kf_lines:
                ma_text += "Key findings:\n" + "\n".join(kf_lines) + "\n"
            if mn_lines:
                ma_text += "Manufactured narratives:\n" + "\n".join(mn_lines) + "\n"
            if actor_lines:
                ma_text += "Actors of concern:\n" + "\n".join(actor_lines) + "\n"
            if cui_text:
                ma_text += "Cui Bono synthesis:\n" + cui_text
            ma_text += f"Intelligence gaps: {_truncate(gaps, 300)}\n"
    else:
        ma_text = "MISSION ANALYST REPORTS: none in last 24h.\n"

    # ── S2 injection evidence (grouped by analyst) ──
    s2_by_analyst: dict[str, list] = {}
    for inj in injections:
        a = inj.get("analyst") or "?"
        s2_by_analyst.setdefault(a, []).append(inj)

    s2_text = "S2 INJECTION FINDINGS (grouped by analyst, exhaustive):\n"
    for analyst in ("S2-A", "S2-B", "S2-C", "S2-D", "S2-E", "S2-GAP"):
        rows = s2_by_analyst.get(analyst, [])
        if not rows:
            s2_text += f"\n{analyst}: no findings.\n"
            continue
        s2_text += f"\n{analyst}: {len(rows)} finding(s)\n"
        for inj in rows[:8]:  # cap per-analyst at 8 to keep tokens sane
            ev = _safe_jsonb(inj.get("evidence")) or {}
            ph = _safe_jsonb(inj.get("flagged_phrases")) or []
            conf = inj.get("confidence_score") or 0.0
            itype = inj.get("injection_type") or "?"
            desc = ""
            if isinstance(ev, dict):
                desc = (
                    ev.get("description")
                    or ev.get("primary_narrative")
                    or ev.get("analyst_note")
                    or ev.get("q1")
                    or ""
                )
            phrase_str = ""
            if isinstance(ph, list) and ph:
                phrase_str = " | ".join(str(p) for p in ph[:5] if p)
            s2_text += f"  [{itype} conf={conf:.2f}] {_truncate(str(desc), 220)}\n"
            if phrase_str:
                s2_text += f"    phrases: {_truncate(phrase_str, 200)}\n"

    # ── S1 lens reports ──
    s1_text = "S1 LENS REPORTS (canary observations):\n"
    if lens_reports:
        for lr in lens_reports[:8]:
            s1_text += (
                f"\n[{lr.get('cycle')}] {lr.get('domain_focus')} "
                f"(quality {lr.get('quality_score')})\n"
                f"  Summary: {_truncate(lr.get('summary'), 350)}\n"
            )
            fft = lr.get("food_for_thought")
            if fft:
                s1_text += f"  Food for thought signal: {_truncate(fft, 200)}\n"
    else:
        s1_text += "(no S1 reports in window)\n"

    # ── S3 (may be stale or absent) ──
    s3_text = "S3 LONG-HORIZON INTELLIGENCE (may be stale; use with caution):\n"
    if s3_latest:
        for pos, row in sorted(s3_latest.items()):
            s3_text += (
                f"\n{pos} [{row.get('report_type')}] generated_at={row.get('generated_at')}\n"
                f"  Summary: {_truncate(row.get('summary'), 400)}\n"
            )
            fd = row.get("first_domino")
            if fd:
                s3_text += f"  First domino: {_truncate(fd, 200)}\n"
    else:
        s3_text += "(no S3 reports available — long-horizon context limited)\n"

    # ── Reference pool (Path B) ──
    ref_text = (
        f"\nREFERENCE POOL — articles collected today ({today_str}, "
        f"total {len(references)}):\n"
        "When you cite an article in Parts 1 or 2, write the REF ID inline like "
        "[REF-YYYYMMDD-NNNN]. Only cite REF IDs from THIS list; do not invent IDs. "
        "If no specific article supports a claim, state the claim without a citation.\n"
    )
    if references:
        # Cap to keep prompt sane; titles are short. ~734 refs at avg 80 chars = ~60k chars
        # which is fine within Opus 4.7's 1M context, but pricier. Cap to 500 most recent.
        for ref in references[:500]:
            ref_text += (
                f"{ref.get('ref_id')} | {ref.get('source_name')} | "
                f"{_truncate(ref.get('title'), 140)}\n"
            )
        if len(references) > 500:
            ref_text += f"... and {len(references) - 500} more not shown to limit prompt size.\n"
    else:
        ref_text += "(no articles available — write Parts 1-2 without citations)\n"

    # ── Master prompt ──
    prompt = f"""You are the lead investigator producing the daily Project Lens Forensic Research Report.

Today's date: {today_str}

This report is read by GCSP (Geneva Centre for Security Policy) educators and emerging global leaders studying information warfare. It is a forensic investigation, not a news summary. It works backwards from observed evidence to identify what shaped today's information environment.

CRITICAL FORMATTING RULES — FOLLOW EXACTLY:
1. NO markdown syntax. No #, no *, no _, no --- dividers, no backticks.
2. Section headers in ALL CAPS followed by colon and newline.
3. Sub-headers in Title Case followed by colon and newline.
4. Bullets start with "- " (dash space).
5. Plain analytical prose. Justified paragraphs. No decoration.

CRITICAL CONTENT RULES:
6. NO predictions anywhere. Not in Part 3, not in Part 4, not anywhere. State only what evidence shows happened or is happening.
7. Cite REF IDs inline in Parts 1-2 using [REF-YYYYMMDD-NNNN] format from the reference pool. Do not fabricate IDs.
8. Part 3 (Food for Thought) is investigative-journalism narrative tone — questions a reader should sit with, not answers the analyst is asserting.
9. Total report length: approximately 2500 words across all four parts.

EVIDENCE INPUT FOLLOWS — analyze it forensically:

{ma_text}

{s2_text}

{s1_text}

{s3_text}

{ref_text}

WRITE EXACTLY THIS STRUCTURE (no other sections, no preamble):

PART 1 — DETECTION:

Be exhaustive. List every distinct injection pattern, manufactured narrative, coordination signal, emotional manipulation method, and legitimacy gap that the evidence shows operating today. For each pattern, state precisely what was detected, which analyst position observed it, the confidence level, and which articles in the reference pool support the finding (cite REF IDs). Do not summarize — enumerate. The goal is a complete forensic inventory of how the information environment was shaped in the last 24 hours.

Subsections required (in this order):

Injection Patterns Detected:
List every injection_type observed across S2-A, the actors involved, and the cognitive effect produced. Cite REF IDs.

Coordination Signals:
State whether S2-B detected coordination. If detection failed (e.g., quota), state that. Describe what cross-source patterns S1's cross-lens analysis flagged.

Emotional Manipulation Methods:
For each S2-C finding, name the emotion targeted, the manipulation score, and the specific 5-step sequence detected.

Adversary Narratives:
State each adversary narrative S2-D identified, with consistency score and key claims.

Legitimacy Gaps:
List every actor S2-E flagged as low-legitimacy, with the gap signal it produces.

Broken Window Findings:
State exactly what S2-GAP identified as missing from mainstream coverage today, with evidence of what is being underreported.

PART 2 — RECOVERY:

What did Project Lens itself recover from today's evidence that mainstream news did not surface? This is the value-added analysis layer. Be specific about which findings would not have emerged from passive news consumption.

Subsections required:

Synthesized Findings From Cross-Position Analysis:
What does the combination of S2-A + S2-D + S2-E reveal that no single position would show? What does the Mission Analyst's cui bono synthesis identify as the convergent beneficiary, and what evidence chain supports it? Cite REF IDs.

Structural Patterns Behind Surface Events:
Reference S3-A (7-day patterns) and S3-D (30-day structural trends) where data is available. State what is forming structurally beneath today's loud events. If S3 data is stale, state the staleness explicitly and constrain claims to what S2 evidence alone supports.

Architect Hypothesis:
Looking at the convergent pattern across S2-A injections, S2-D adversary narratives, S2-E legitimacy gaps, and the Mission Analyst's cui bono synthesis — identify candidate upstream architects by evidence weight. An "architect" is the state apparatus or non-state network whose strategic interests, capability, and historical pattern are most consistent with shaping today's information environment, not merely the most visible executor. Name candidates by specific office or apparatus (for example, "Xi Office / CCP Politburo," "Putin Office / Kremlin apparatus," "Trump Office / US executive branch," "Khamenei Office / IRGC," "Netanyahu Office," "Modi Office"). Do not use country names as shorthand for apparatus — the Chinese people, Russian people, American people, Iranian people, and all peoples are distinct from whatever apparatus currently governs them, and are Project Lens's intended beneficiaries, not its targets. When the source pool contains state-media from a given actor, name that apparatus as a candidate for evaluation even if same-day S2 signatures are absent — absence of injection signature over a short window does not mean absence of architect role. For each candidate, construct the evidence chain from today's articles (cite REF IDs) and note relevant S3-B historical analog when available. Evaluate alternative explanations for each candidate (coordinated influence vs editorial-market pressure vs organic concern alignment vs sample contamination). Weight legitimacy asymmetrically per the project's baseline: elected, term-bounded, constitutionally-constrained offices are analyzed within the scope of their mandate and term; unelected-indefinite or fake-elected offices are analyzed without the benefit of a bounded-mandate assumption. This is forensic accuracy about what legitimacy means, not political preference. Avoid single-villain reduction — oppression of peoples typically emerges from convergent architectures (state apparatus networks, oligarchic capture, surveillance-tech supply chains, financial-secrecy systems, transnational organized crime overlapping with state actors), not lone villains. If the evidence chain is too thin on this cycle to name any architect with confidence, say so explicitly and constrain the finding to "no architect attribution is supported by today's evidence alone."

Intelligence Gaps Worth Naming:
State explicitly what could not be determined from today's evidence. What sources were absent or returned no data. What claims would require evidence not currently in the pool.

PART 3 — FOOD FOR THOUGHT:

Switch to investigative-journalism narrative voice. This is questions for the reader to sit with, written in flowing prose, not a list of bullets. Three to five paragraphs. No predictions — only questions about what the evidence pattern means. Examples of acceptable framing: "What does it tell us when X and Y converge in the same week?" "Why might Z be invisible in mainstream coverage despite its scale?" "Whose interests are served when the loud event drowns out the quiet one?"

Do not assert answers. Do not predict outcomes. Pose the questions and let the reader carry them.

PART 4 — REFERENCES:

Leave this section empty — just the header "PART 4 — REFERENCES:". The references will be auto-appended after rendering, listing every REF ID you cited in Parts 1 and 2.

Begin the report now. Do not add a preamble or title — start directly with "PART 1 — DETECTION:".
"""

    log.info(f"Prompt built: {len(prompt)} chars (~{len(prompt)//4} tokens estimated)")
    return prompt


# ── Anthropic call ────────────────────────────────────────────────────────────

def call_opus(client, prompt: str) -> str | None:
    """Call Opus 4.7 with retry on transient failure."""
    for attempt in (1, 2, 3):
        try:
            log.info(f"Calling {MODEL} attempt {attempt}/3 (max_tokens={MAX_TOKENS})...")
            t0 = time.time()
            msg = client.messages.create(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                messages=[{"role": "user", "content": prompt}],
            )
            elapsed = time.time() - t0
            text = msg.content[0].text
            usage = getattr(msg, "usage", None)
            input_t = getattr(usage, "input_tokens", "?") if usage else "?"
            output_t = getattr(usage, "output_tokens", "?") if usage else "?"
            cost_in = (input_t / 1_000_000 * 5.0) if isinstance(input_t, int) else 0
            cost_out = (output_t / 1_000_000 * 25.0) if isinstance(output_t, int) else 0
            log.info(
                f"Opus done: {len(text)} chars / ~{len(text.split())} words, "
                f"in={input_t} out={output_t}, cost~${cost_in + cost_out:.4f}, {elapsed:.1f}s"
            )
            return text
        except Exception as e:
            log.error(f"Opus attempt {attempt} failed: {e}")
            if attempt < 3:
                time.sleep(20)
    return None


# ── Citation validator (Path B safety net) ────────────────────────────────────

REF_ID_PATTERN = re.compile(r"\[REF-\d{8}-\d{3,5}\]")


def validate_citations(report_text: str, valid_ref_ids: set) -> tuple[str, dict]:
    """Strip any [REF-...] from the report that isn't in the valid pool.

    Returns (cleaned_text, stats_dict) where stats_dict has counts.
    Per LR-080: this is the post-write integrity check on LLM output.
    """
    cited = REF_ID_PATTERN.findall(report_text)
    cited_clean = [c.strip("[]") for c in cited]
    cited_set = set(cited_clean)
    valid = cited_set & valid_ref_ids
    invalid = cited_set - valid_ref_ids

    cleaned = report_text
    if invalid:
        for bad in invalid:
            # Replace "[REF-bad-id]" with empty + collapse double spaces
            cleaned = cleaned.replace(f"[{bad}]", "")
        cleaned = re.sub(r"  +", " ", cleaned)
        cleaned = re.sub(r" \.", ".", cleaned)
        cleaned = re.sub(r" ,", ",", cleaned)

    stats = {
        "total_citations_attempted": len(cited),
        "unique_citations_attempted": len(cited_set),
        "valid_citations": len(valid),
        "invalid_stripped": len(invalid),
        "invalid_ids": sorted(invalid),
        "valid_ids": sorted(valid),
    }
    log.info(
        f"Citation validation: attempted={len(cited_set)} valid={len(valid)} stripped={len(invalid)}"
    )
    if invalid:
        log.warning(f"Stripped fabricated REF IDs: {sorted(invalid)[:5]}{'...' if len(invalid) > 5 else ''}")
    return cleaned, stats


# ── DOCX renderer (python-docx, per LR-067) ───────────────────────────────────

def render_docx(report_text: str, valid_cited_refs: list, references: list,
                output_path: str, report_date: str) -> bool:
    """Render report as DOCX with auto-appended Part 4 references list."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as e:
        log.error(f"python-docx import failed: {e}")
        return False

    try:
        doc = Document()
        sec = doc.sections[0]
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)

        # Page-number footer
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        for tag, ftype in (("w:fldChar", "begin"), ("w:instrText", None), ("w:fldChar", "end")):
            el = OxmlElement(tag)
            if ftype:
                el.set(qn("w:fldCharType"), ftype)
            else:
                el.text = "PAGE"
            run._r.append(el)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Title
        tp = doc.add_paragraph()
        tr = tp.add_run(f"Project Lens — Forensic Research Report\n{report_date}")
        tr.bold = True
        tr.font.size = Pt(14)
        tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        pPr = tp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), "1a1a2e")
        pBdr.append(bot)
        pPr.append(pBdr)
        doc.add_paragraph()

        # Body — split lines, classify each
        for line in report_text.split("\n"):
            s = line.strip()
            if not s:
                doc.add_paragraph()
                continue

            stripped_no_colon = s.rstrip(":")
            normalized = stripped_no_colon.replace(" ", "").replace("—", "").replace("-", "")
            is_section = (
                s.endswith(":")
                and len(s.split()) >= 2
                and normalized.isupper()
                and len(s) < 80
            )
            is_bullet = s.startswith("- ")
            is_sub = (
                s.endswith(":")
                and len(s) < 80
                and not is_section
                and s[0].isupper()
            )

            if is_section:
                p = doc.add_paragraph()
                r = p.add_run(s.rstrip(":"))
                r.bold = True
                r.font.size = Pt(13)
                r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                pP = p._p.get_or_add_pPr()
                pB = OxmlElement("w:pBdr")
                b2 = OxmlElement("w:bottom")
                b2.set(qn("w:val"), "single")
                b2.set(qn("w:sz"), "2")
                b2.set(qn("w:space"), "4")
                b2.set(qn("w:color"), "cccccc")
                pB.append(b2)
                pP.append(pB)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
            elif is_sub:
                p = doc.add_paragraph()
                r = p.add_run(s)
                r.bold = True
                r.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
            elif is_bullet:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(s[2:]).font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph(s)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(4)
                for r in p.runs:
                    r.font.size = Pt(10)

        # Auto-append Part 4 references list (only those actually cited & validated)
        if valid_cited_refs:
            ref_lookup = {r["ref_id"]: r for r in references if r.get("ref_id")}
            doc.add_paragraph()
            for ref_id in sorted(valid_cited_refs):
                ref = ref_lookup.get(ref_id)
                if not ref:
                    continue
                p = doc.add_paragraph()
                r1 = p.add_run(f"{ref_id}  ")
                r1.bold = True
                r1.font.size = Pt(9)
                p.add_run(f"{ref.get('source_name', '')} — {ref.get('title', '')}").font.size = Pt(9)
                if ref.get("url"):
                    p.add_run(f"\n  {ref.get('url')}").font.size = Pt(8)
                p.paragraph_format.space_after = Pt(3)

        doc.save(output_path)
        log.info(f"DOCX rendered: {output_path}")
        return True
    except Exception as e:
        log.error(f"DOCX render failed: {e}")
        return False


# ── Telegram delivery ─────────────────────────────────────────────────────────

def build_caption(report_text: str, citation_stats: dict, report_date: str,
                  cap: int = TELEGRAM_CAPTION_CAP) -> str:
    """Build the file caption per LENS-016 budget: header 80 / Detection 240 /
    Recovery 240 / Questions 330 / footer 60."""
    # Extract first ~2 sentences of each part as a teaser (best-effort)
    def section_teaser(text: str, marker: str, char_budget: int) -> str:
        idx = text.find(marker)
        if idx == -1:
            return ""
        rest = text[idx + len(marker):]
        end = rest.find("PART ")
        if end == -1:
            end = len(rest)
        body = rest[:end].strip()
        body = REF_ID_PATTERN.sub("", body).strip()
        return _truncate(body, char_budget)

    detection = section_teaser(report_text, "PART 1 — DETECTION:", 240)
    recovery = section_teaser(report_text, "PART 2 — RECOVERY:", 240)
    questions = section_teaser(report_text, "PART 3 — FOOD FOR THOUGHT:", 330)

    header = f"📑 Project Lens Forensic Report — {report_date}"
    footer = f"Refs cited: {citation_stats.get('valid_citations', 0)}"

    parts = [header, "", detection, "", recovery, "", questions, "", footer]
    caption = "\n".join(p for p in parts if p)
    return _truncate(caption, cap)


def send_telegram(docx_path: str, caption: str, filename: str) -> bool:
    import requests
    token = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        log.warning("Telegram credentials missing — skip send")
        return False
    try:
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        with open(docx_path, "rb") as f:
            resp = requests.post(
                url,
                data={"chat_id": chat_id, "caption": caption},
                files={"document": (filename, f,
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=60,
            )
        ok = resp.status_code == 200
        log.info(f"Telegram: {'OK' if ok else f'FAILED {resp.status_code} {resp.text[:200]}'}")
        return ok
    except Exception as e:
        log.error(f"Telegram send failed: {e}")
        return False


# ── Main orchestration ───────────────────────────────────────────────────────

def run_forensic_report(dry_run: bool = False) -> dict:
    start = time.time()
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    log.info(f"=== FORENSIC REPORT START | {today} | dry_run={dry_run} ===")

    # ── Phase 1: Connect to Supabase only (Anthropic deferred until needed) ──
    try:
        sb = get_supabase()
    except Exception as e:
        log.error(f"Supabase init failed: {e}")
        return {"status": "ERROR", "phase": "supabase_init", "error": str(e)}

    # ── Phase 2: Fetch evidence ──
    macros = fetch_macro_reports(sb)
    injections = fetch_injection_reports(sb)
    lens_reports = fetch_lens_reports(sb)
    s3_latest = fetch_system3_latest(sb)
    references = fetch_reference_pool(sb)

    if not macros and not injections:
        log.warning("No MA and no injection data in window — skipping report.")
        return {"status": "SKIP", "reason": "no_evidence_in_window"}

    # ── Phase 3: Build prompt ──
    prompt = build_prompt(macros, injections, lens_reports, s3_latest, references)

    # ── Phase 4: Dry-run early exit ──
    if dry_run:
        dry_path = f"/tmp/forensic_prompt_{today}.txt"
        try:
            with open(dry_path, "w", encoding="utf-8") as f:
                f.write(prompt)
            log.info(f"[DRY-RUN] Prompt written to {dry_path}")
        except Exception as e:
            log.warning(f"[DRY-RUN] Could not write prompt to disk: {e}")
        return {
            "status": "DRY_RUN_OK",
            "prompt_chars": len(prompt),
            "prompt_estimated_tokens": len(prompt) // 4,
            "estimated_cost_usd": round((len(prompt) // 4 / 1_000_000) * 5.0
                                        + (MAX_TOKENS / 1_000_000) * 25.0, 4),
            "macros_used": len(macros),
            "injections_used": len(injections),
            "lens_reports_used": len(lens_reports),
            "s3_positions_with_data": sorted(s3_latest.keys()),
            "references_in_pool": len(references),
            "prompt_path": dry_path,
        }

    # ── Phase 5: Anthropic API call (real) ──
    try:
        client = get_anthropic()
    except Exception as e:
        log.error(f"Anthropic init failed: {e}")
        return {"status": "ERROR", "phase": "anthropic_init", "error": str(e)}

    raw_report = call_opus(client, prompt)
    if not raw_report:
        return {"status": "ERROR", "phase": "opus_call", "error": "all 3 attempts failed"}

    # ── Phase 6: Validate citations (Path B safety net) ──
    valid_ref_id_set = {r["ref_id"] for r in references if r.get("ref_id")}
    cleaned_report, citation_stats = validate_citations(raw_report, valid_ref_id_set)
    valid_cited = citation_stats["valid_ids"]

    # ── Phase 7: Render DOCX ──
    filename = f"{today}_ProjectLens_Forensic_DC2200.docx"
    docx_path = os.path.join(tempfile.gettempdir(), filename)
    docx_ok = render_docx(cleaned_report, valid_cited, references, docx_path, today)
    if not docx_ok:
        return {"status": "ERROR", "phase": "render_docx",
                "raw_report_preview": cleaned_report[:500]}

    # ── Phase 8: Telegram delivery ──
    caption = build_caption(cleaned_report, citation_stats, today)
    sent = send_telegram(docx_path, caption, filename)

    elapsed = round(time.time() - start, 1)
    log.info(f"=== FORENSIC REPORT DONE | sent={sent} | {elapsed}s ===")
    return {
        "status": "OK",
        "elapsed_s": elapsed,
        "report_words": len(cleaned_report.split()),
        "citations": citation_stats,
        "docx_path": docx_path,
        "telegram_sent": sent,
    }


# ── CLI ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Project Lens Forensic Research Report")
    parser.add_argument("--dry-run", action="store_true",
                        help="Fetch data + build prompt, skip API/render/send (cost: $0)")
    args = parser.parse_args()

    try:
        from dotenv import load_dotenv
        load_dotenv()
    except ImportError:
        pass

    result = run_forensic_report(dry_run=args.dry_run)
    print(json.dumps(result, indent=2, default=str))
    sys.exit(0 if result.get("status") in ("OK", "DRY_RUN_OK", "SKIP") else 1)
