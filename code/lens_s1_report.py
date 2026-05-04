"""
lens_s1_report.py — S1 Canary Intelligence Report
Project Lens | LENS-023
Model: mistral-small-latest (free)
Purpose: Full quality docx of what the 4-lens canary detected today.
         Companion to S2 report — operator compares S1 (raw signal)
         vs S2 (how that signal was shaped) to reveal manipulation delta.
Output: YYYYMMDD_S1_Canary_Intelligence_DC{time}.docx → Telegram
"""

import os, json, time, logging, tempfile, requests
from datetime import datetime, timezone, timedelta
from typing import Optional

logging.basicConfig(level=logging.INFO,
    format="%(asctime)s [S1-RPT] %(levelname)s %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("s1_report")

MODEL       = "mistral-small-latest"
TEMPERATURE = 0.3
MAX_TOKENS  = 4000
TELEGRAM_CAPTION_CAP = 950


# ── Data fetch ────────────────────────────────────────────────────────────────

def fetch_s1_data(run_id: Optional[str] = None) -> dict:
    from supabase import create_client
    sb = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])

    # S1 lens reports — latest 4
    s1 = sb.table("lens_reports") \
        .select("id,domain_focus,summary,quality_score,cycle,generated_at,articles_used") \
        .order("generated_at", desc=True).limit(4).execute().data or []

    # Cross-lens signals from most recent run
    cross = []
    if s1:
        latest_run = sb.table("lens_reports") \
            .select("domain_focus,summary,quality_score") \
            .order("generated_at", desc=True).limit(8).execute().data or []
        cross = latest_run

    # Collection stats — last 12h
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=12)).isoformat()
    arts = sb.table("lens_raw_articles") \
        .select("source_name,domain,collected_at") \
        .gte("collected_at", cutoff).execute().data or []

    # Source breakdown
    sources = {}
    domains = {}
    for a in arts:
        s = a.get("source_name", "Unknown")
        d = a.get("domain", "UNKNOWN")
        sources[s] = sources.get(s, 0) + 1
        domains[d] = domains.get(d, 0) + 1

    # Top entity
    top_entity = (sb.table("lens_entities") \
        .select("canonical_name,total_mentions,entity_type") \
        .order("total_mentions", desc=True).limit(5).execute().data or [])

    # 7-day threat trend
    cutoff_7d = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
    trend = sb.table("lens_macro_reports") \
        .select("threat_level,created_at") \
        .gte("created_at", cutoff_7d) \
        .order("created_at", desc=True).limit(7).execute().data or []

    log.info(f"S1 report: {len(s1)} lenses, {len(arts)} articles, {len(top_entity)} entities")
    return {
        "s1": s1, "articles": arts, "sources": sources,
        "domains": domains, "entities": top_entity, "trend": trend
    }


# ── AI prompt ─────────────────────────────────────────────────────────────────

def build_s1_prompt(data: dict) -> str:
    s1 = data["s1"]
    sources = data["sources"]
    domains = data["domains"]
    entities = data["entities"]
    trend = data["trend"]
    total_arts = len(data["articles"])

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    avg_q = round(sum(r.get("quality_score") or 0 for r in s1) / len(s1), 1) if s1 else 0

    prompt = f"""You are the S1 Canary Intelligence Analyst for Project Lens, an OSINT geopolitical intelligence system built for GCSP educators.

Today is {date_str}. You have received the raw signal outputs from the 4-lens canary system. Your task is to write a comprehensive, analytical intelligence report that reads like a professional geopolitical briefing — long analytical paragraphs, cause-effect reasoning, no bullet points in the body text.

=== S1 CANARY RAW DATA ===

COLLECTION POOL: {total_arts} articles collected in the last 12 hours
Domain breakdown: {json.dumps(domains)}
Top sources by volume: {json.dumps(dict(sorted(sources.items(), key=lambda x: -x[1])[:10]))}
Average lens quality: {avg_q}/10

LENS REPORTS (4 lenses):
"""
    for r in s1:
        prompt += f"""
Lens: {r.get('domain_focus', 'UNKNOWN')} | Quality: {r.get('quality_score', 0)}/10 | Articles used: {r.get('articles_used', '?')}
Summary: {r.get('summary', 'No summary')}
"""

    if entities:
        prompt += "\nENTITY REGISTRY (most mentioned actors today):\n"
        for e in entities:
            prompt += f"  {e.get('canonical_name')} ({e.get('entity_type')}) — {e.get('total_mentions')} mentions\n"

    if trend:
        levels = [t.get('threat_level', '?') for t in trend[:5]]
        prompt += f"\n7-DAY THREAT TREND: {' → '.join(levels)}\n"

    prompt += """

=== REPORT INSTRUCTIONS ===

Write a full intelligence report using exactly this structure. Each Part must be written as flowing analytical prose — multiple long paragraphs, not bullet points. Minimum 150 words per Part.

PART A — THE COLLECTION LANDSCAPE
Analyze today's article collection: what domains are overrepresented or underrepresented? What does the source distribution reveal about the information environment? Which sources dominated and why does that matter?

PART B — WHAT THE FOUR LENSES DETECTED
For each lens, synthesize the key signal it detected today. Explain what each lens contributes and how the four perspectives relate to each other. Where do they agree? Where do they diverge?

PART C — CROSS-LENS CONVERGENCE AND DIVERGENCE
What patterns emerge when you compare all four lenses simultaneously? Which narratives were confirmed across multiple lenses (high confidence signals)? Which appeared in only one lens (isolated or potentially planted)?

PART D — THE ENTITY PICTURE
Who are the most active actors in today's information environment? What does their prominence reveal about today's geopolitical narrative structure?

PART E — THE CANARY'S VERDICT
What is the overall picture the canary is painting today? Write a 2-3 paragraph analytical verdict that a GCSP educator could read as the definitive S1 raw intelligence summary. What should the reader take to System 2 for deeper analysis?

Write in formal intelligence briefing style. Never use bullet points in body paragraphs. Use PART A, PART B format exactly for headings."""

    return prompt


# ── AI call ───────────────────────────────────────────────────────────────────

def call_mistral(prompt: str) -> Optional[str]:
    api_key = os.environ.get("MISTRAL_API_KEY", "")
    if not api_key:
        log.error("MISTRAL_API_KEY not set")
        return None

    for attempt in range(1, 4):
        try:
            log.info(f"S1 report calling Mistral (attempt {attempt})")
            r = requests.post(
                "https://api.mistral.ai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={
                    "model": MODEL,
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": MAX_TOKENS,
                    "temperature": TEMPERATURE,
                },
                timeout=120
            )
            if r.status_code == 200:
                text = r.json()["choices"][0]["message"]["content"].strip()
                log.info(f"S1 report: {len(text)} chars generated")
                return text
            log.warning(f"Mistral {r.status_code} attempt {attempt}")
            time.sleep(20 * attempt)
        except Exception as e:
            log.error(f"Mistral call failed attempt {attempt}: {e}")
            time.sleep(15)
    return None


# ── Docx renderer ─────────────────────────────────────────────────────────────

def render_docx(report_text: str, date_str: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROJECT LENS — S1 CANARY INTELLIGENCE REPORT")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"{date_str}  |  System 1  |  4-Lens Canary  |  Mistral-small")
    sr.font.size = Pt(10); sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disclaimer.add_run("RAW SIGNAL — Compare with S2 Report to identify manipulation delta")
    dr.italic = True; dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x99, 0x44, 0x00)

    doc.add_paragraph()

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph(); continue
        if line.startswith("PART ") and "—" in line:
            doc.add_heading(line, level=1)
        elif line.isupper() and len(line) < 80 and ":" not in line:
            doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    time_str = datetime.now(timezone.utc).strftime("%H%M")
    fname = f"{date_str.replace('-', '')}_S1_Canary_Intelligence_DC{time_str}.docx"
    tmp_path = os.path.join(tempfile.gettempdir(), fname)
    doc.save(tmp_path)
    log.info(f"S1 docx saved: {tmp_path}")
    return tmp_path


# ── Telegram ──────────────────────────────────────────────────────────────────

def send_telegram_doc(docx_path: str, caption: str) -> bool:
    token   = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        log.warning("Telegram keys not set — skipping"); return False
    try:
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        with open(docx_path, "rb") as f:
            r = requests.post(url, data={"chat_id": chat_id, "caption": caption,
                "parse_mode": "HTML"}, files={"document": (
                    os.path.basename(docx_path), f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )}, timeout=60)
        if r.status_code == 200:
            log.info("S1 docx sent to Telegram"); return True
        log.error(f"Telegram sendDocument failed: {r.status_code} {r.text[:200]}"); return False
    except Exception as e:
        log.error(f"Telegram send failed: {e}"); return False


def send_telegram_text(text: str) -> bool:
    token   = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        return False
    try:
        r = requests.post(
            f"https://api.telegram.org/bot{token}/sendMessage",
            json={"chat_id": chat_id, "text": text, "parse_mode": "HTML"}, timeout=10)
        return r.status_code == 200
    except Exception:
        return False


# ── Entry point ───────────────────────────────────────────────────────────────

def run_s1_report(run_id: Optional[str] = None) -> dict:
    start = time.time()
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    log.info(f"=== S1 CANARY REPORT START | {date_str} ===")

    try:
        data = fetch_s1_data(run_id)
    except Exception as e:
        log.error(f"Data fetch failed: {e}")
        return {"status": "ERROR", "error": str(e)}

    if not data["s1"]:
        log.warning("No S1 lens reports found")
        return {"status": "NO_DATA"}

    prompt = build_s1_prompt(data)
    report_text = call_mistral(prompt)

    if not report_text:
        log.error("Mistral failed to generate S1 report")
        return {"status": "AI_FAILED"}

    try:
        docx_path = render_docx(report_text, date_str)
    except Exception as e:
        log.error(f"Docx render failed: {e}")
        return {"status": "DOCX_FAILED", "error": str(e)}

    # Short Telegram intro
    s1 = data["s1"]
    avg_q = round(sum(r.get("quality_score") or 0 for r in s1) / len(s1), 1) if s1 else 0
    total_arts = len(data["articles"])
    intro = (
        f"📡 <b>S1 Canary Intelligence Report — {date_str}</b>\n"
        f"4 lenses | {total_arts} articles | Avg quality {avg_q}/10\n"
        f"<i>Full analytical report attached — compare with S2 to see manipulation delta</i>"
    )
    send_telegram_text(intro)
    time.sleep(1)

    # Full docx caption
    caption = (
        f"📡 S1 Canary Intelligence Report — {date_str}\n"
        f"System 1 | 4-Lens Canary | {total_arts} articles analyzed\n"
        f"Avg quality: {avg_q}/10\n"
        f"Parts: Collection Landscape | Lens Findings | Convergence | Entities | Verdict"
    )[:TELEGRAM_CAPTION_CAP]

    sent = send_telegram_doc(docx_path, caption)
    elapsed = round(time.time() - start, 1)

    log.info(f"=== S1 REPORT COMPLETE | sent={sent} | {elapsed}s ===")
    return {"status": "COMPLETE" if sent else "SEND_FAILED", "elapsed": elapsed}


if __name__ == "__main__":
    from dotenv import load_dotenv; load_dotenv()
    result = run_s1_report()
    print(result)
