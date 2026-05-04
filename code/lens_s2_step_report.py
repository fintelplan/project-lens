"""
lens_s2_step_report.py — S2 Information Shaping Intelligence Report
Project Lens | LENS-023
Model: mistral-small-latest (free)
Purpose: Full quality docx of how today's information environment was shaped.
         Shows injection patterns, adversary narratives, coordination signals,
         emotional architecture, legitimacy gaps, and what S1 missed.
Output: YYYYMMDD_S2_Shaping_Intelligence_DC{time}.docx → Telegram
"""

import os, json, time, logging, tempfile, requests
from datetime import datetime, timezone, timedelta
from typing import Optional

logging.basicConfig(level=logging.INFO,
    format="%(asctime)s [S2-RPT] %(levelname)s %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("s2_report")

MODEL       = "mistral-small-latest"
TEMPERATURE = 0.3
MAX_TOKENS  = 4500
TELEGRAM_CAPTION_CAP = 950


# ── Data fetch ────────────────────────────────────────────────────────────────

def fetch_s2_data(run_id: Optional[str] = None) -> dict:
    from supabase import create_client
    sb = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])

    # Get latest run_id if not provided
    if not run_id:
        ma = sb.table("lens_macro_reports") \
            .select("run_id,threat_level,executive_summary,quality_score,contamination_depth,created_at") \
            .order("created_at", desc=True).limit(1).execute().data or []
        if ma:
            run_id = ma[0].get("run_id", "")
            ma_data = ma[0]
        else:
            ma_data = {}
    else:
        ma = sb.table("lens_macro_reports") \
            .select("run_id,threat_level,executive_summary,quality_score,contamination_depth,created_at") \
            .eq("run_id", run_id).execute().data or []
        ma_data = ma[0] if ma else {}

    # All injection reports for this run
    inj = sb.table("injection_reports") \
        .select("analyst,injection_type,confidence_score,flagged_phrases,evidence,created_at") \
        .eq("run_id", run_id) \
        .order("created_at", desc=True).limit(50).execute().data or []

    # S1 lens reports for context
    s1 = sb.table("lens_reports") \
        .select("domain_focus,summary,quality_score") \
        .order("generated_at", desc=True).limit(4).execute().data or []

    # S2-F operation detections
    cutoff_24h = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
    s2f = sb.table("lens_operation_detections") \
        .select("state_actor_lens,operation_count,confidence,scored_at") \
        .gte("scored_at", cutoff_24h) \
        .eq("not_applicable", False) \
        .order("confidence", desc=True).limit(10).execute().data or []

    log.info(f"S2 report: run_id={run_id}, {len(inj)} injection reports, {len(s2f)} S2-F detections")
    return {"run_id": run_id, "ma": ma_data, "inj": inj, "s1": s1, "s2f": s2f}


# ── AI prompt ─────────────────────────────────────────────────────────────────

def ev(raw):
    if not raw: return {}
    if isinstance(raw, dict): return raw
    try: return json.loads(raw)
    except: return {}

def build_s2_prompt(data: dict) -> str:
    run_id  = data["run_id"]
    ma      = data["ma"]
    inj     = data["inj"]
    s1      = data["s1"]
    s2f     = data["s2f"]
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # Parse injection reports by analyst
    s2a_rows = [i for i in inj if i.get("analyst") == "S2-A"]
    s2b_rows = [i for i in inj if i.get("analyst") == "S2-B"]
    s2c_rows = [i for i in inj if i.get("analyst") == "S2-C"]
    s2d_rows = [i for i in inj if i.get("analyst") == "S2-D"]
    s2e_rows = [i for i in inj if i.get("analyst") == "S2-E"]
    gap_rows = [i for i in inj if i.get("analyst") == "S2-GAP"]

    prompt = f"""You are the S2 Information Shaping Analyst for Project Lens, an OSINT geopolitical intelligence system built for GCSP educators.

Today is {date_str} | Run ID: {run_id}

Your task is to write a comprehensive analytical intelligence report on HOW TODAY'S INFORMATION ENVIRONMENT WAS SHAPED. This is a professional intelligence briefing — long analytical paragraphs, cause-effect reasoning, no bullet points in body text. Each Part minimum 200 words.

=== MISSION ANALYST VERDICT ===
Threat Level: {ma.get('threat_level', 'UNKNOWN')}
Quality Score: {ma.get('quality_score', 0)}
Contamination Depth: {ma.get('contamination_depth', 'UNKNOWN')}
Executive Summary: {ma.get('executive_summary', 'Not available')}

=== S1 CANARY RAW SIGNAL (what the uncontaminated canary detected) ===
"""
    for r in s1:
        prompt += f"Lens {r.get('domain_focus','?')} (quality {r.get('quality_score',0)}/10): {r.get('summary','')}\n\n"

    prompt += "\n=== S2-A INJECTION TRACER (narrative contamination patterns) ===\n"
    for row in s2a_rows[:4]:
        e = ev(row.get("evidence"))
        phrases = row.get("flagged_phrases") or []
        if isinstance(phrases, str):
            try: phrases = json.loads(phrases)
            except: phrases = [phrases]
        prompt += f"Injection type: {row.get('injection_type')} | Confidence: {row.get('confidence_score',0):.0%}\n"
        desc = e.get("description","") or e.get("q1","") or e.get("raw","")
        if desc: prompt += f"Analysis: {desc[:400]}\n"
        if phrases: prompt += f"Flagged phrases: {', '.join(str(p) for p in phrases[:5])}\n"
        prompt += "\n"

    prompt += "\n=== S2-B COORDINATION ANALYZER (cross-source coordination signals) ===\n"
    for row in s2b_rows[:3]:
        e = ev(row.get("evidence"))
        itype = row.get("injection_type","")
        if itype and itype not in ("NO_COORDINATION","NONE",""):
            prompt += f"Coordination type: {itype} | Confidence: {row.get('confidence_score',0):.0%}\n"
            detail = e.get("description","") or e.get("dominant_narrative","") or e.get("analyst_note","")
            if detail: prompt += f"Detail: {detail[:400]}\n"
            prompt += "\n"

    prompt += "\n=== S2-C EMOTION DECODER (emotional manipulation architecture) ===\n"
    for row in s2c_rows[:4]:
        e = ev(row.get("evidence"))
        emotion = e.get("dominant_emotion","") or e.get("q3","") or e.get("frame","")
        score = row.get("confidence_score",0)
        steps = e.get("steps","") or e.get("manipulation_steps","")
        prompt += f"Emotional target: {emotion} | Manipulation score: {score:.0%}\n"
        if steps: prompt += f"Steps detected: {str(steps)[:300]}\n"
        prompt += "\n"

    prompt += "\n=== S2-D ADVERSARY NARRATIVE (state actor messaging) ===\n"
    for row in s2d_rows[:2]:
        e = ev(row.get("evidence"))
        nar = e.get("primary_narrative","") or e.get("q1","") or ""
        claims = e.get("key_claims",[]) or []
        actors = e.get("named_actors",{}) or {}
        tone = e.get("emotional_tone","")
        counter = e.get("counter_narrative_target","")
        prompt += f"Primary adversarial narrative: {nar[:500]}\n"
        if actors: prompt += f"Named actors — Heroes: {actors.get('heroes',[])} | Villains: {actors.get('villains',[])} | Victims: {actors.get('victims',[])}\n"
        if tone: prompt += f"Emotional tone: {tone}\n"
        if counter: prompt += f"Counter-narrative target: {counter[:200]}\n"
        if claims:
            prompt += f"Key claims ({len(claims)} total):\n"
            for c in claims[:4]:
                if isinstance(c, dict):
                    prompt += f"  - {c.get('claim','')[:150]} [source: {c.get('source','')}]\n"
        prompt += "\n"

    prompt += "\n=== S2-E LEGITIMACY FILTER (who lacks democratic mandate) ===\n"
    for row in s2e_rows[:4]:
        e = ev(row.get("evidence"))
        verdict = e.get("verdict","") or e.get("legitimacy_verdict","") or e.get("q6","")
        if verdict: prompt += f"Legitimacy assessment: {verdict[:400]}\n\n"

    prompt += "\n=== S2-GAP ANALYSIS (what S1 missed — Broken Window) ===\n"
    for row in gap_rows[:2]:
        e = ev(row.get("evidence"))
        key = e.get("key_gap_finding","")
        missed = e.get("missed_by_s1",[]) or []
        severity = e.get("severity","")
        if key: prompt += f"Key gap: {key[:500]}\nSeverity: {severity}\n"
        if isinstance(missed, list):
            for m in missed[:3]:
                if isinstance(m, dict):
                    prompt += f"Missed story: {m.get('story','')[:200]}\nWhy it matters: {m.get('why_significant','')[:200]}\n\n"

    if s2f:
        prompt += "\n=== S2-F OPERATION DETECTIONS (structural information operations) ===\n"
        for d in s2f[:5]:
            prompt += f"Actor: {d.get('state_actor_lens')} | Operations detected: {d.get('operation_count')} | Confidence: {d.get('confidence',0):.0%}\n"

    prompt += """

=== REPORT INSTRUCTIONS ===

Write a full intelligence report using exactly this structure. Each Part must be written as flowing analytical prose with multiple long paragraphs. Minimum 200 words per Part. No bullet points in body text.

PART A — THE INJECTION ARCHITECTURE
Analyze the full pattern of narrative injection detected today. What types of contamination were used? What does the combination of injection methods reveal about the intent? How do the flagged phrases function as cognitive triggers? Explain the mechanism of manipulation with analytical depth.

PART B — THE ADVERSARY NARRATIVE MAP
What are the state actors communicating today? What is the primary adversarial message? Who are the heroes and villains in their framing? What Western/opposing narrative are they specifically countering? What behavior or belief are they trying to engineer?

PART C — COORDINATION AND EMOTIONAL ARCHITECTURE
Was today's information environment coordinated across sources? What emotional states were being deliberately engineered — urgency, distrust, fear, division? How do the coordination signals and emotional manipulation work together as an integrated system?

PART D — THE LEGITIMACY GAP AND BROKEN WINDOW
Which actors lack democratic legitimacy but dominated today's narrative? What significant stories did the mainstream canary (S1) miss entirely? Why are those absences analytically significant? What does the gap between what was covered and what was ignored reveal?

PART E — S2-F STRUCTURAL OPERATIONS
What information operations were detected at the structural level (S2-F)? Which state actor lenses show the highest operational activity? What does the pattern of operations suggest about strategic intent?

PART F — MISSION ANALYST SYNTHESIS
Synthesize all S2 findings into a comprehensive verdict. What is the overall contamination picture? How should GCSP educators adjust their reading of today's media environment based on these findings?

Write in formal intelligence briefing style. Use PART A, PART B format exactly for headings."""

    return prompt


# ── AI call ───────────────────────────────────────────────────────────────────

def call_mistral(prompt: str) -> Optional[str]:
    api_key = os.environ.get("MISTRAL_API_KEY", "")
    if not api_key:
        log.error("MISTRAL_API_KEY not set"); return None
    for attempt in range(1, 4):
        try:
            log.info(f"S2 report calling Mistral (attempt {attempt})")
            r = requests.post(
                "https://api.mistral.ai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": MODEL, "messages": [{"role": "user", "content": prompt}],
                      "max_tokens": MAX_TOKENS, "temperature": TEMPERATURE},
                timeout=120)
            if r.status_code == 200:
                text = r.json()["choices"][0]["message"]["content"].strip()
                log.info(f"S2 report: {len(text)} chars generated")
                return text
            log.warning(f"Mistral {r.status_code} attempt {attempt}")
            time.sleep(20 * attempt)
        except Exception as e:
            log.error(f"Mistral call failed: {e}"); time.sleep(15)
    return None


# ── Docx renderer ─────────────────────────────────────────────────────────────

def render_docx(report_text: str, date_str: str, ma: dict) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROJECT LENS — S2 INFORMATION SHAPING INTELLIGENCE REPORT")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    threat = ma.get("threat_level", "UNKNOWN")
    sr = sub.add_run(f"{date_str}  |  System 2  |  THREAT: {threat}  |  Mistral-small")
    sr.font.size = Pt(10); sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    depth = ma.get("contamination_depth", "UNKNOWN")
    dep = doc.add_paragraph()
    dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dep.add_run(f"Contamination Depth: {depth}  |  Quality: {ma.get('quality_score',0):.2f}")
    dr.italic = True; dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph()

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph(); continue
        if line.startswith("PART ") and "—" in line:
            doc.add_heading(line, level=1)
        elif line.isupper() and len(line) < 80 and ":" not in line:
            doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    time_str = datetime.now(timezone.utc).strftime("%H%M")
    fname = f"{date_str.replace('-','')}_S2_Shaping_Intelligence_DC{time_str}.docx"
    tmp_path = os.path.join(tempfile.gettempdir(), fname)
    doc.save(tmp_path)
    log.info(f"S2 docx saved: {tmp_path}")
    return tmp_path


# ── Telegram ──────────────────────────────────────────────────────────────────

def send_telegram_doc(docx_path: str, caption: str) -> bool:
    token   = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        log.warning("Telegram keys not set"); return False
    try:
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        with open(docx_path, "rb") as f:
            r = requests.post(url, data={"chat_id": chat_id, "caption": caption,
                "parse_mode": "HTML"}, files={"document": (
                    os.path.basename(docx_path), f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )}, timeout=60)
        if r.status_code == 200:
            log.info("S2 docx sent to Telegram"); return True
        log.error(f"Telegram error: {r.status_code}"); return False
    except Exception as e:
        log.error(f"Telegram send failed: {e}"); return False


def send_telegram_text(text: str) -> bool:
    token   = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id: return False
    try:
        r = requests.post(f"https://api.telegram.org/bot{token}/sendMessage",
            json={"chat_id": chat_id, "text": text, "parse_mode": "HTML"}, timeout=10)
        return r.status_code == 200
    except: return False


# ── Entry point ───────────────────────────────────────────────────────────────

def run_s2_report(run_id: Optional[str] = None) -> dict:
    start = time.time()
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    log.info(f"=== S2 SHAPING REPORT START | {date_str} ===")

    try:
        data = fetch_s2_data(run_id)
    except Exception as e:
        log.error(f"Data fetch failed: {e}")
        return {"status": "ERROR", "error": str(e)}

    if not data["inj"]:
        log.warning("No injection reports found for this run")
        return {"status": "NO_DATA"}

    prompt = build_s2_prompt(data)
    report_text = call_mistral(prompt)
    if not report_text:
        return {"status": "AI_FAILED"}

    try:
        docx_path = render_docx(report_text, date_str, data["ma"])
    except Exception as e:
        log.error(f"Docx render failed: {e}")
        return {"status": "DOCX_FAILED", "error": str(e)}

    ma = data["ma"]
    threat = ma.get("threat_level", "UNKNOWN")
    depth  = ma.get("contamination_depth", "UNKNOWN")

    intro = (
        f"🔬 <b>S2 Information Shaping Report — {date_str}</b>\n"
        f"Threat: <b>{threat}</b> | Contamination: {depth}\n"
        f"<i>Full analytical report attached — 6 parts covering injection patterns, "
        f"adversary narratives, coordination signals, and what S1 missed</i>"
    )
    send_telegram_text(intro)
    time.sleep(1)

    caption = (
        f"🔬 S2 Information Shaping Intelligence Report — {date_str}\n"
        f"Threat: {threat} | Contamination: {depth}\n"
        f"Parts: Injection Architecture | Adversary Narrative | Coordination | "
        f"Legitimacy Gap | S2-F Operations | Mission Analyst Synthesis"
    )[:TELEGRAM_CAPTION_CAP]

    sent = send_telegram_doc(docx_path, caption)
    elapsed = round(time.time() - start, 1)
    log.info(f"=== S2 REPORT COMPLETE | sent={sent} | {elapsed}s ===")
    return {"status": "COMPLETE" if sent else "SEND_FAILED", "elapsed": elapsed}


if __name__ == "__main__":
    from dotenv import load_dotenv; load_dotenv()
    result = run_s2_report()
    print(result)
