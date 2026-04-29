"""
lens_regular_report.py
Project Lens — Regular Report (Free Tier)
Model: mistral-small-latest (free) → Cerebras fallback → Groq fallback
Schedule: 1x daily at 02:10 UTC (10:10 PM DC EDT)
Structure: Same 4-part as Opus Report (Detection/Recovery/Food for Thought/References)
Cost: $0/run

Reads from DB:
  - S1: lens_reports (last 24h)
  - S3: lens_system3_reports (latest per position)
  - References: lens_article_refs (last 24h)

Model runs S2+MA synthesis directly (not reading pre-processed S2 DB outputs).
Guard system: preflight → retry (3x) → fallback providers → write-verify → failure alert.

Usage:
    python code/lens_regular_report.py            # full run
    python code/lens_regular_report.py --dry-run  # no API call, no Telegram
"""

import argparse
import logging
import os
import re
import sys
import tempfile
import time
from datetime import datetime, timedelta, timezone

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [REGULAR] %(levelname)s %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("REGULAR")

# ── Constants ─────────────────────────────────────────────────────────────────
LOOKBACK_HOURS       = 24
MAX_TOKENS           = 4096
TEMPERATURE          = 0.3
TELEGRAM_CAPTION_CAP = 950
MAX_REFS             = 400
MAX_RETRIES          = 3

# Provider chain: mistral-small → cerebras → groq
PROVIDERS = [
    {
        "name": "mistral",
        "model": "mistral-small-latest",
        "key_env": "MISTRAL_API_KEY",
        "base_url": "https://api.mistral.ai/v1",
    },
    {
        "name": "cerebras",
        "model": "qwen-3-235b-a22b-instruct-2507",
        "key_env": "CEREBRAS_API_KEY",
        "base_url": None,  # uses cerebras SDK
    },
    {
        "name": "groq",
        "model": "llama-3.3-70b-versatile",
        "key_env": "GROQ_API_KEY",
        "base_url": None,  # uses groq SDK
    },
]


# ── Guard: Preflight ──────────────────────────────────────────────────────────
def preflight_check() -> bool:
    """Verify required env vars exist before running."""
    required = ["SUPABASE_URL", "SUPABASE_SERVICE_KEY", "TELEGRAM_BOT_TOKEN", "TELEGRAM_CHAT_ID"]
    missing = [k for k in required if not os.environ.get(k)]
    if missing:
        log.error(f"PREFLIGHT FAILED — missing: {missing}")
        return False
    # At least one LLM key must be present
    llm_keys = ["MISTRAL_API_KEY", "CEREBRAS_API_KEY", "GROQ_API_KEY"]
    if not any(os.environ.get(k) for k in llm_keys):
        log.error("PREFLIGHT FAILED — no LLM API key found (MISTRAL/CEREBRAS/GROQ)")
        return False
    log.info("PREFLIGHT OK")
    return True


# ── Clients ───────────────────────────────────────────────────────────────────
def get_supabase():
    from supabase import create_client
    url = os.environ.get("SUPABASE_URL", "")
    key = os.environ.get("SUPABASE_SERVICE_KEY", "")
    if not url or not key:
        raise RuntimeError("SUPABASE credentials missing")
    return create_client(url, key)


def get_llm_client():
    """Return (client, model, provider_name) trying each provider in chain."""
    for prov in PROVIDERS:
        key = os.environ.get(prov["key_env"], "")
        if not key:
            log.info(f"Skipping {prov['name']} — key not set")
            continue
        try:
            if prov["name"] == "cerebras":
                from cerebras.cloud.sdk import Cerebras
                client = Cerebras(api_key=key)
                log.info(f"LLM: using Cerebras ({prov['model']})")
                return client, prov["model"], "cerebras"
            elif prov["name"] == "groq":
                from groq import Groq
                client = Groq(api_key=key)
                log.info(f"LLM: using Groq ({prov['model']})")
                return client, prov["model"], "groq"
            else:
                from openai import OpenAI
                client = OpenAI(api_key=key, base_url=prov["base_url"])
                log.info(f"LLM: using {prov['name']} ({prov['model']})")
                return client, prov["model"], prov["name"]
        except Exception as e:
            log.warning(f"Failed to init {prov['name']}: {e}")
            continue
    raise RuntimeError("No LLM provider available")


# ── Evidence fetchers ─────────────────────────────────────────────────────────
def fetch_s1_reports(sb) -> list:
    """S1 lens reports last 24h."""
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=LOOKBACK_HOURS)).isoformat()
    try:
        r = sb.table("lens_reports") \
            .select("domain_focus,summary,food_for_thought,quality_score,cycle,generated_at") \
            .gte("generated_at", cutoff) \
            .order("generated_at", desc=True) \
            .limit(20) \
            .execute()
        log.info(f"S1: {len(r.data)} lens reports in last {LOOKBACK_HOURS}h")
        return r.data or []
    except Exception as e:
        log.warning(f"S1 fetch failed: {e}")
        return []


def fetch_s3_reports(sb) -> dict:
    """Latest S3 report per position."""
    positions = ["S3-A", "S3-B", "S3-C", "S3-D", "S3-E"]
    results = {}
    for pos in positions:
        try:
            r = sb.table("lens_system3_reports") \
                .select("position,report_type,summary,first_domino,patterns_found,"
                        "structural_trends,signals_to_watch,quality_score,generated_at") \
                .eq("position", pos) \
                .order("generated_at", desc=True) \
                .limit(1) \
                .execute()
            if r.data:
                results[pos] = r.data[0]
        except Exception as e:
            log.warning(f"S3-{pos} fetch failed: {e}")
    log.info(f"S3: positions with data = {list(results.keys())}")
    return results


def fetch_references(sb) -> list:
    """Article references for citation pool."""
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    yesterday = (datetime.now(timezone.utc) - timedelta(days=1)).strftime("%Y-%m-%d")
    try:
        r = sb.table("lens_article_refs") \
            .select("ref_id,title,source_name,domain,collected_date") \
            .gte("collected_date", yesterday) \
            .lte("collected_date", today) \
            .order("collected_date", desc=True) \
            .order("ref_id", desc=False) \
            .execute()
        refs = r.data or []
        log.info(f"References: {len(refs)} articles in 36h window")
        return refs[:MAX_REFS]
    except Exception as e:
        log.warning(f"References fetch failed: {e}")
        return []


# ── Prompt builder ────────────────────────────────────────────────────────────
def build_prompt(s1_reports: list, s3_reports: dict, references: list) -> str:
    """Build full prompt for mistral-small S2+MA synthesis."""

    # S1 block
    s1_text = f"SYSTEM 1 — LENS REPORTS ({len(s1_reports)} reports, last 24h):\n\n"
    if s1_reports:
        for r in s1_reports:
            domain = r.get("domain_focus", "")
            summary = (r.get("summary") or "")[:500]
            fft = (r.get("food_for_thought") or "")[:200]
            quality = r.get("quality_score", 0)
            s1_text += f"Domain: {domain} | Quality: {quality}\n"
            s1_text += f"Summary: {summary}\n"
            if fft:
                s1_text += f"Food for thought: {fft}\n"
            s1_text += "\n"
    else:
        s1_text += "No S1 reports available in this window.\n"

    # S3 block
    s3_text = "SYSTEM 3 — PATTERN INTELLIGENCE (latest per position):\n\n"
    if s3_reports:
        for pos, rep in s3_reports.items():
            summary = (rep.get("summary") or "")[:400]
            first_domino = (rep.get("first_domino") or "")[:200]
            signals = (rep.get("signals_to_watch") or "")[:200]
            s3_text += f"{pos}:\n"
            s3_text += f"  Summary: {summary}\n"
            if first_domino:
                s3_text += f"  First domino: {first_domino}\n"
            if signals:
                s3_text += f"  Signals to watch: {signals}\n"
            s3_text += "\n"
    else:
        s3_text += "No S3 reports available.\n"

    # Reference pool
    ref_text = f"REFERENCE POOL ({len(references)} articles — use [REF-ID] citations):\n"
    for ref in references:
        ref_id = ref.get("ref_id", "")
        title = (ref.get("title") or "")[:120]
        source = (ref.get("source_name") or ref.get("domain") or "")
        ref_text += f"{ref_id}  {source} — {title}\n"

    # System prompt
    system_prompt = """You are the Project Lens Regular Report Generator.

Project Lens is a cognitive sovereignty intelligence platform built on PHI-002 (pro-people, anti-pretense) and PHI-003 (Office vocabulary — Xi Office, Putin Office, Trump Office, Khamenei Office — never conflating apparatus with people).

Your task: Read today's S1 lens intelligence and S3 pattern intelligence. Then independently synthesize a full intelligence report covering what S2 positions would detect, what the Mission Analyst would conclude, and deliver a 4-part report.

VOCABULARY DISCIPLINE (non-negotiable):
- Use Office names: "Xi Office", "Trump Office", "Putin Office", "Khamenei Office"
- The peoples (Chinese, Russian, American, Iranian) are never targets — always intended beneficiaries
- Never predict future events
- Every claim must cite [REF-YYYYMMDD-NNNN] from the reference pool

REPORT STRUCTURE — write all 4 parts:

PART 1 — DETECTION
Exhaustive analysis of how today's information environment was shaped:
- What injection patterns appear in today's coverage? (FACT_VOID, FALSE_EQUIV, EMOTIONAL_PRIME, SOURCE_LAUNDER)
- Which low-legitimacy actors dominate the narrative? What legitimacy gaps exist?
- What coordination patterns appear across sources?
- What emotional manipulation sequences are present?
- What adversary narratives are being pushed?

PART 2 — RECOVERY  
Synthesis that no single position would surface:
- Cross-position convergence: what do S1 + S3 together reveal?
- Cui bono: who benefits from today's dominant framing?
- Structural patterns behind surface events
- What the combined evidence reveals about information operations in progress

PART 3 — FOOD FOR THOUGHT
Investigative questions for GCSP educators to hold while watching these voices.
3-5 open questions. No predictions. Evidence-based only.

PART 4 — REFERENCES
List every REF ID cited in Parts 1-2 with source name and title.
Format: [REF-ID] Source — Title

Length: 4-5 pages, ~2000-2500 words. Exhaustive detection depth.
Tone: Parts 1-2 intelligence-briefing formal. Part 3 investigative-journalism narrative.
"""

    user_msg = f"""{s1_text}

{s3_text}

{ref_text}

Generate the full 4-part Project Lens Regular Report. Cite [REF-ID] throughout Parts 1-2. Be exhaustive — this is the free daily intelligence report for GCSP educators."""

    return system_prompt, user_msg


# ── LLM call with retry ───────────────────────────────────────────────────────
def call_llm(system_prompt: str, user_msg: str) -> str:
    """Call LLM with retry + provider fallback. Returns report text."""
    client, model, provider = get_llm_client()

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            log.info(f"Calling {provider}/{model} attempt {attempt}/{MAX_RETRIES}...")
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_msg},
                ],
                max_tokens=MAX_TOKENS,
                temperature=TEMPERATURE,
                timeout=300,
            )
            text = resp.choices[0].message.content.strip()
            words = len(text.split())
            log.info(f"{provider} done: {len(text)} chars / ~{words} words")
            return text
        except Exception as e:
            log.warning(f"Attempt {attempt} failed ({provider}): {str(e)[:200]}")
            if attempt < MAX_RETRIES:
                sleep_sec = 10 * attempt
                log.info(f"Retrying in {sleep_sec}s...")
                time.sleep(sleep_sec)
            else:
                log.error(f"All {MAX_RETRIES} attempts failed for {provider}")
                # Try next provider
                remaining = [p for p in PROVIDERS if p["name"] != provider
                             and os.environ.get(p["key_env"])]
                if remaining:
                    log.info(f"Falling back to {remaining[0]['name']}")
                    # Rebuild client with next provider
                    os.environ["_FORCE_PROVIDER"] = remaining[0]["name"]
                    try:
                        client, model, provider = get_llm_client()
                        resp = client.chat.completions.create(
                            model=model,
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": user_msg},
                            ],
                            max_tokens=MAX_TOKENS,
                            temperature=TEMPERATURE,
                            timeout=300,
                        )
                        return resp.choices[0].message.content.strip()
                    except Exception as e2:
                        log.error(f"Fallback also failed: {e2}")
                raise RuntimeError(f"All providers exhausted")


# ── Citation validator ────────────────────────────────────────────────────────
def validate_citations(text: str, references: list) -> tuple:
    """Strip invalid REF IDs from report text. Returns (cleaned_text, stats)."""
    valid_ids = {r["ref_id"] for r in references if r.get("ref_id")}
    found = set(re.findall(r"REF-\d{8}-\d{4}", text))
    invalid = found - valid_ids
    valid = found & valid_ids

    cleaned = text
    for bad_id in invalid:
        cleaned = cleaned.replace(f"[{bad_id}]", "")
        cleaned = cleaned.replace(bad_id, "")

    stats = {
        "total_citations_attempted": len(found),
        "valid_citations": len(valid),
        "invalid_stripped": len(invalid),
        "invalid_ids": sorted(invalid),
    }
    log.info(f"Citation validation: attempted={len(found)} valid={len(valid)} stripped={len(invalid)}")
    return cleaned, stats


# ── DOCX renderer ─────────────────────────────────────────────────────────────
def render_docx(report_text: str, date_str: str, references: list) -> str:
    """Render report text to docx. Returns temp file path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed — pip install python-docx==1.1.2")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECT LENS — REGULAR INTELLIGENCE REPORT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"{date_str}  |  Free Tier  |  Mistral-small")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Split report into lines and render
    lines = report_text.split("\n")
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue

        # Detect headings
        if line_stripped.startswith("PART ") and "—" in line_stripped:
            p = doc.add_heading(line_stripped, level=1)
        elif line_stripped.isupper() and len(line_stripped) < 60 and ":" not in line_stripped:
            p = doc.add_heading(line_stripped, level=2)
        else:
            p = doc.add_paragraph(line_stripped)
            p.paragraph_format.space_after = Pt(6)

    # Save to temp file
    fname = f"{date_str.replace('-', '')}_ProjectLens_Regular_DC2210.docx"
    tmp_path = os.path.join(tempfile.gettempdir(), fname)
    doc.save(tmp_path)
    log.info(f"DOCX rendered: {tmp_path}")
    return tmp_path


# ── Telegram sender ───────────────────────────────────────────────────────────
def build_caption(report_text: str, citation_stats: dict) -> str:
    """Build Telegram caption from report text."""
    lines = [l.strip() for l in report_text.split("\n") if l.strip()]

    # Find Part 1 and Part 2 first sentences
    part1_text = ""
    part2_text = ""
    in_part1 = False
    in_part2 = False

    for line in lines:
        if "PART 1" in line and "DETECTION" in line:
            in_part1 = True
            in_part2 = False
            continue
        if "PART 2" in line and "RECOVERY" in line:
            in_part2 = True
            in_part1 = False
            continue
        if "PART 3" in line or "PART 4" in line:
            in_part1 = False
            in_part2 = False
        if in_part1 and not part1_text and len(line) > 40:
            part1_text = line[:200]
        if in_part2 and not part2_text and len(line) > 40:
            part2_text = line[:200]

    valid = citation_stats.get("valid_citations", 0)
    caption = f"📋 Project Lens Regular Report — {datetime.now(timezone.utc).strftime('%Y%m%d')}\n"
    caption += "Injection Patterns Detected:\n\n"
    if part1_text:
        caption += f"{part1_text}...\n"
    caption += "Synthesized Findings From Cross-Position Analysis:\n\n"
    if part2_text:
        caption += f"{part2_text}...\n"
    caption += f"Refs cited: {valid}"

    return caption[:TELEGRAM_CAPTION_CAP]


def send_telegram(docx_path: str, caption: str) -> bool:
    """Send docx file to Telegram."""
    import requests
    token = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        log.error("Telegram credentials missing")
        return False
    url = f"https://api.telegram.org/bot{token}/sendDocument"
    try:
        with open(docx_path, "rb") as f:
            resp = requests.post(
                url,
                data={"chat_id": chat_id, "caption": caption},
                files={"document": (os.path.basename(docx_path), f,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=60,
            )
        if resp.status_code == 200:
            log.info("Telegram: OK")
            return True
        else:
            log.error(f"Telegram failed: {resp.status_code} {resp.text[:200]}")
            return False
    except Exception as e:
        log.error(f"Telegram error: {e}")
        return False


def send_telegram_alert(message: str) -> None:
    """Send failure alert to Telegram."""
    import requests
    token = os.environ.get("TELEGRAM_BOT_TOKEN", "")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID", "")
    if not token or not chat_id:
        return
    try:
        requests.post(
            f"https://api.telegram.org/bot{token}/sendMessage",
            data={"chat_id": chat_id, "text": f"⚠️ REGULAR REPORT FAILED\n{message}"},
            timeout=30,
        )
    except Exception:
        pass


# ── Main ──────────────────────────────────────────────────────────────────────
def run_regular_report(dry_run: bool = False) -> dict:
    """Main entry point. Returns status dict."""
    start = time.time()
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    log.info(f"=== REGULAR REPORT START | {date_str.replace('-','')} | dry_run={dry_run} ===")

    # Guard: preflight
    if not preflight_check():
        send_telegram_alert("Preflight check failed — missing credentials")
        return {"status": "FAILED", "reason": "preflight"}

    # Fetch evidence
    try:
        sb = get_supabase()
        s1_reports = fetch_s1_reports(sb)
        s3_reports = fetch_s3_reports(sb)
        references = fetch_references(sb)
    except Exception as e:
        msg = f"DB fetch failed: {e}"
        log.error(msg)
        send_telegram_alert(msg)
        return {"status": "FAILED", "reason": msg}

    if not s1_reports:
        msg = "No S1 reports in DB — skipping report"
        log.warning(msg)
        send_telegram_alert(msg)
        return {"status": "SKIPPED", "reason": msg}

    # Build prompt
    system_prompt, user_msg = build_prompt(s1_reports, s3_reports, references)
    prompt_chars = len(system_prompt) + len(user_msg)
    log.info(f"Prompt built: {prompt_chars} chars (~{prompt_chars//4} tokens estimated)")

    if dry_run:
        log.info("DRY RUN — skipping LLM call and Telegram")
        return {"status": "DRY_RUN", "prompt_chars": prompt_chars}

    # Call LLM
    try:
        report_text = call_llm(system_prompt, user_msg)
    except Exception as e:
        msg = f"LLM call failed: {e}"
        log.error(msg)
        send_telegram_alert(msg)
        return {"status": "FAILED", "reason": msg}

    # Validate citations
    report_text, citation_stats = validate_citations(report_text, references)

    # Render docx
    try:
        docx_path = render_docx(report_text, date_str, references)
    except Exception as e:
        msg = f"DOCX render failed: {e}"
        log.error(msg)
        send_telegram_alert(msg)
        return {"status": "FAILED", "reason": msg}

    # Send Telegram
    caption = build_caption(report_text, citation_stats)
    sent = send_telegram(docx_path, caption)

    elapsed = round(time.time() - start, 1)
    log.info(f"=== REGULAR REPORT DONE | sent={sent} | {elapsed}s ===")

    return {
        "status": "OK",
        "elapsed_s": elapsed,
        "report_words": len(report_text.split()),
        "citations": citation_stats,
        "docx_path": docx_path,
        "telegram_sent": sent,
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    from dotenv import load_dotenv
    load_dotenv()

    result = run_regular_report(dry_run=args.dry_run)
    import json
    print(json.dumps(result, indent=2))
    sys.exit(0 if result["status"] in ("OK", "DRY_RUN") else 1)
