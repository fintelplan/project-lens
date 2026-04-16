"""
lens_sonnet_report.py
Project Lens — Sonnet 4.6 Intelligence Report

5 pages: S2 + S3 in detail
Title: date/time only (e.g. "April 17, 2026 10:28 PM")
No cover, no intro, no footer note
Page numbers only

Runs 2x/day:
  15:28 UTC = 11:00 PM DC = 10:28 AM Thailand (wait, DC 11pm = 03:28 UTC)
  Actually:
  15:28 UTC = 11:28 AM DC = 10:28 PM Thailand
  03:28 UTC = 11:28 PM DC = 10:28 AM Thailand
"""

import os, json, logging, time, tempfile, subprocess
from datetime import datetime, timezone, timedelta
from typing import Optional

logging.basicConfig(level=logging.INFO,
    format="%(asctime)s [SONNET-RPT] %(levelname)s %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("SONNET_RPT")

MODEL      = "claude-sonnet-4-6"
MAX_TOKENS = 8000


def get_supabase():
    from supabase import create_client
    return create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_SERVICE_KEY"])

def get_anthropic():
    import anthropic
    return anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])


# ── Reuse fetch functions ─────────────────────────────────────────────────────

def fetch_s2_full(sb) -> list:
    """Fetch all S2 injection reports from last 6 hours."""
    try:
        cutoff = (datetime.now(timezone.utc) - timedelta(hours=6)).isoformat()
        r = sb.table("injection_reports") \
            .select("analyst,injection_type,confidence_score,flagged_phrases,evidence,created_at") \
            .gte("created_at", cutoff) \
            .order("created_at", desc=True).limit(20).execute()
        data = r.data or []
        log.info(f"S2: {len(data)} injection reports loaded")
        return data
    except Exception as e:
        log.warning(f"S2 fetch failed: {e}")
        return []

def fetch_s3a(sb) -> dict:
    try:
        r = sb.table("lens_system3_reports") \
            .select("summary,first_domino,patterns_found,quality_score,generated_at") \
            .eq("position", "S3-A") \
            .order("generated_at", desc=True).limit(1).execute()
        return r.data[0] if r.data else {}
    except Exception as e:
        log.warning(f"S3-A fetch failed: {e}")
        return {}

def fetch_s3d(sb) -> dict:
    try:
        r = sb.table("lens_system3_reports") \
            .select("summary,first_domino,quality_score,generated_at") \
            .eq("position", "S3-D") \
            .order("generated_at", desc=True).limit(1).execute()
        return r.data[0] if r.data else {}
    except Exception as e:
        log.warning(f"S3-D fetch failed: {e}")
        return {}

def fetch_ma(sb) -> dict:
    try:
        r = sb.table("lens_macro_reports") \
            .select("threat_level,executive_summary,cui_bono_synthesis,quality_score,created_at") \
            .order("created_at", desc=True).limit(1).execute()
        return r.data[0] if r.data else {}
    except Exception as e:
        log.warning(f"MA fetch failed: {e}")
        return {}


# ── Prompt ────────────────────────────────────────────────────────────────────

def build_prompt(s2, s3a, s3d, ma) -> str:
    def ev(raw):
        if not raw: return {}
        if isinstance(raw, dict): return raw
        try: return json.loads(raw)
        except: return {}

    def phrases(raw):
        if not raw: return []
        if isinstance(raw, list): return raw
        try: return json.loads(raw)
        except: return [str(raw)]

    # Format S2 by analyst
    analysts = {}
    for inj in s2:
        a = inj.get("analyst","?")
        if a not in analysts:
            analysts[a] = inj

    s2_text = "S2 DATA:\n"
    for analyst in ["S2-A","S2-B","S2-C","S2-D","S2-E","S2-GAP"]:
        inj = analysts.get(analyst)
        if not inj:
            s2_text += f"\n{analyst}: No data this cycle.\n"
            continue
        itype = inj.get("injection_type","?")
        conf  = inj.get("confidence_score",0) or 0
        e     = ev(inj.get("evidence"))
        ph    = phrases(inj.get("flagged_phrases"))[:4]
        desc  = e.get("description","") or e.get("primary_narrative","") or e.get("q1","") or e.get("raw","")
        s2_text += f"\n{analyst}: {itype} — {conf:.0%}\n"
        if desc: s2_text += f"  {str(desc)[:400]}\n"
        if ph:   s2_text += f"  Phrases: {' | '.join(str(p) for p in ph if p)}\n"

    # MA Cui Bono
    cui = ev(ma.get("cui_bono_synthesis") or {}) if ma else {}
    ma_text = ""
    if ma:
        ma_text = f"\nMA REPORT:\nThreat: {ma.get('threat_level','?')} | Quality: {ma.get('quality_score',0):.2f}\n"
        ma_text += f"{(ma.get('executive_summary') or '')[:500]}\n"
        if isinstance(cui, dict):
            ma_text += f"Cui Bono primary: {cui.get('primary_beneficiary','?')}\n"
            ma_text += f"Convergence: {cui.get('convergence','?')}\n"
            ma_text += f"Evidence: {cui.get('evidence','?')}\n"

    # S3
    s3a_text = ""
    if s3a:
        s3a_text = f"\nS3-A (quality {s3a.get('quality_score',0):.2f}):\n"
        s3a_text += (s3a.get("summary","") or "")[:800] + "\n"
        dom = s3a.get("first_domino","")
        if dom: s3a_text += f"First Domino: {dom}\n"
        s3a_text += "ACH and sectarian trap: not yet in DB schema\n"

    s3d_text = ""
    if s3d:
        s3d_text = f"\nS3-D (quality {s3d.get('quality_score',0):.2f}):\n"
        s3d_text += (s3d.get("summary","") or "")[:500] + "\n"
        dom = s3d.get("first_domino","")
        if dom: s3d_text += f"Structural First Domino: {dom}\n"

    prompt = f"""You are writing an intelligence report. Write exactly 5 pages of precise, factual analysis.

CRITICAL FORMATTING RULES:
1. No markdown. No hashtags, asterisks, dashes as dividers.
2. Section headers in ALL CAPS followed by colon.
3. Sub-headers in Title Case followed by colon.
4. Bullet items start with dash and space (- item).
5. Plain prose paragraphs. Nothing decorative.
6. Be precise and complete. No padding. No summaries.

DATA:
{s2_text}
{ma_text}
{s3a_text}
{s3d_text}

WRITE EXACTLY THIS STRUCTURE:

SYSTEM 2 — INJECTION ANALYSIS:

S2-A Injection Detection:
State the injection type, confidence level, and what it means in practice.
List the flagged phrases. Explain what cognitive effect they are designed to produce.

S2-B Coordination Analysis:
State whether coordination was detected. If yes: what pattern, which sources, what timing.
If no: state clearly and explain what was observed instead.

S2-C Emotional Framing:
State the dominant emotion being manufactured. State the manipulation score.
Explain the 5-step emotion sequence being deployed and its intended outcome.

SYSTEM 2 — ADVERSARY AND GAP:

S2-D Adversary Narrative:
State the primary adversary narrative exactly. State consistency score.
List the key claims being pushed. Identify target audiences.

S2-E Legitimacy Assessment:
List every actor flagged. State their legitimacy tier.
State the legitimacy gap signal — what this means for the information environment.

S2-GAP Broken Window:
State what S1 missed. List stories that were absent from mainstream analysis.
State the key gap finding. Explain why this absence matters.

CUI BONO — WHO BENEFITS:

Primary Beneficiary:
State who benefits most from today's complete S2 injection pattern.
Explain which S2 positions all point to the same actor.

Secondary and Tertiary Beneficiaries:
State secondary beneficiary and why.
State tertiary beneficiary and why.

Convergence Assessment:
State convergence level: CONFIRMED, PROBABLE, or UNCLEAR.
Explain the evidence chain that supports this attribution.

SYSTEM 3 — PATTERN INTELLIGENCE:

7-Day Pattern:
State the pattern forming. Be specific about the evidence base.
State the event sequence — in what order did things happen and what does the order reveal.

Distraction and Quiet Event:
State what loud event is consuming analytical bandwidth.
State what quiet structural event is happening behind it.

Hidden Builder:
State who is gaining structural advantage quietly. What is the evidence.

First Domino:
State exactly what becomes inevitable if current patterns continue.
This is not prediction — it is the logical endpoint of observed evidence.

SYSTEM 3 — CHECKS AND VERDICT:

ACH Check:
State the strongest evidence that contradicts the pattern analysis.
State confidence level and why the dominant reading still holds despite this contradiction.

Sectarian Trap Status:
State whether ethnic, religious, or political tension manufacturing is detected.
If yes: state the group, the amplifier, organic or manufactured.
If no: state clearly.

S3-D Structural:
State the 30-day structural finding. State the structural first domino.

Mission Analyst Verdict:
State the threat level.
State the single most important finding from this cycle.
State 3 signals to watch in the next cycle."""

    return prompt


# ── Call Sonnet ───────────────────────────────────────────────────────────────

def call_sonnet(client, prompt: str) -> Optional[str]:
    for attempt in range(1, 4):
        try:
            log.info(f"Calling {MODEL} attempt {attempt}/3...")
            msg = client.messages.create(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                messages=[{"role": "user", "content": prompt}]
            )
            report = msg.content[0].text
            log.info(f"Report: {len(report.split())} words")
            return report
        except Exception as e:
            log.error(f"Attempt {attempt} failed: {e}")
            if attempt < 3: time.sleep(30)
    return None


# ── DOCX generation ───────────────────────────────────────────────────────────

def generate_docx(report_text: str, output_path: str, title: str) -> bool:
    try:
        import sys, shutil
        is_win = sys.platform.startswith("win")

        # Find node path explicitly
        node_cmd = shutil.which("node") or "node"
        npm_cmd  = shutil.which("npm")  or "npm"

        result = subprocess.run([node_cmd, "--version"],
            capture_output=True, text=True, shell=is_win)
        if result.returncode != 0:
            log.error(f"node not found: {result.stderr}")
            return False

        subprocess.run([npm_cmd, "install", "-g", "docx"],
            capture_output=True, shell=is_win)

        sections = []
        current_header = None
        current_content = []

        for line in report_text.split('\n'):
            stripped = line.strip()
            if not stripped:
                if current_content: current_content.append("")
                continue
            words = stripped.rstrip(':').split()
            is_header = (
                stripped.endswith(':') and
                len(words) >= 2 and
                stripped.rstrip(':').replace(' ','').replace('—','').replace('-','').isupper() and
                len(stripped) < 80
            )
            if is_header:
                if current_header:
                    sections.append({"h": current_header, "c": '\n'.join(current_content)})
                current_header = stripped.rstrip(':')
                current_content = []
            else:
                current_content.append(stripped)
        if current_header:
            sections.append({"h": current_header, "c": '\n'.join(current_content)})

        sections_json = json.dumps(sections, ensure_ascii=False)
        title_safe = title.replace("'","\\'")

        js = f"""
const {{ Document, Packer, Paragraph, TextRun, AlignmentType,
         LevelFormat, BorderStyle, PageNumber, Footer }} = require('docx');
const fs = require('fs');

const sec = {sections_json};
const children = [];

// Title — date/time only
children.push(new Paragraph({{
  spacing: {{ before: 0, after: 480 }},
  border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 6, color: '1a1a2e', space: 4 }} }},
  children: [new TextRun({{ text: '{title_safe}', bold: true, size: 32, font: 'Arial' }})]
}}));

for (const s of sec) {{
  // Section header
  children.push(new Paragraph({{
    spacing: {{ before: 360, after: 160 }},
    border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 1, color: 'cccccc', space: 4 }} }},
    children: [new TextRun({{ text: s.h, bold: true, size: 24, font: 'Arial', color: '1a1a2e' }})]
  }}));

  for (const line of s.c.split('\\n')) {{
    const t = line.trim();
    if (!t) {{ children.push(new Paragraph({{ spacing: {{ after: 60 }} }})); continue; }}

    if (t.startsWith('- ')) {{
      children.push(new Paragraph({{
        numbering: {{ reference: 'b', level: 0 }},
        spacing: {{ after: 60 }},
        children: [new TextRun({{ text: t.slice(2), size: 20, font: 'Arial' }})]
      }}));
      continue;
    }}

    const isSub = t.endsWith(':') && t.length < 70 && t[0] === t[0].toUpperCase() && !t.includes('  ');
    if (isSub) {{
      children.push(new Paragraph({{
        spacing: {{ before: 160, after: 80 }},
        children: [new TextRun({{ text: t, bold: true, size: 21, font: 'Arial' }})]
      }}));
      continue;
    }}

    children.push(new Paragraph({{
      spacing: {{ after: 100 }},
      alignment: AlignmentType.JUSTIFIED,
      children: [new TextRun({{ text: t, size: 20, font: 'Arial' }})]
    }}));
  }}
}}

const doc = new Document({{
  numbering: {{
    config: [{{
      reference: 'b',
      levels: [{{ level: 0, format: LevelFormat.BULLET, text: '•',
        alignment: AlignmentType.LEFT,
        style: {{ paragraph: {{ indent: {{ left: 720, hanging: 360 }} }} }} }}]
    }}]
  }},
  sections: [{{
    properties: {{
      page: {{
        size: {{ width: 11906, height: 16838 }},
        margin: {{ top: 1440, right: 1440, bottom: 1440, left: 1440 }}
      }}
    }},
    footers: {{
      default: new Footer({{
        children: [new Paragraph({{
          alignment: AlignmentType.CENTER,
          children: [new TextRun({{ children: [PageNumber.CURRENT], size: 18, font: 'Arial', color: '888888' }})]
        }})]
      }})
    }},
    children: children
  }}]
}});

Packer.toBuffer(doc).then(buf => {{
  fs.writeFileSync('{output_path}', buf);
  console.log('OK');
}}).catch(e => {{ console.error(e.message); process.exit(1); }});
"""
        import pathlib
        project_dir = str(pathlib.Path(__file__).parent.parent)
        js_path = str(pathlib.Path(project_dir) / "lens_gen_tmp.js")
        with open(js_path, 'w', encoding='utf-8') as f:
            f.write(js)

        result = subprocess.run([node_cmd, js_path], capture_output=True, text=True, timeout=60, shell=is_win, cwd=project_dir)
        if result.returncode == 0 and os.path.exists(output_path):
            os.remove(js_path)
            log.info(f"DOCX generated: {output_path}")
            return True
        log.error(f"Node failed: {result.stderr[:200]}")
        return False
    except Exception as e:
        log.error(f"DOCX failed: {e}")
        return False


# ── Telegram send ─────────────────────────────────────────────────────────────

def send_telegram(docx_path: str, title: str) -> bool:
    import requests
    token   = os.environ.get("TELEGRAM_BOT_TOKEN","")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID","")
    if not token or not chat_id: return False
    try:
        fname = title.replace(' ','_').replace(',','').replace(':','') + ".docx"
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        with open(docx_path,'rb') as f:
            resp = requests.post(url,
                data={"chat_id": chat_id},
                files={"document": (fname, f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=60)
        ok = resp.status_code == 200
        log.info(f"Telegram: {'OK' if ok else 'FAILED'}")
        return ok
    except Exception as e:
        log.error(f"Telegram failed: {e}")
        return False


# ── Main ──────────────────────────────────────────────────────────────────────

def run_sonnet_report() -> dict:
    start = time.time()

    # Thai time for title (UTC+7)
    utc_now = datetime.now(timezone.utc)
    dc_time = utc_now - timedelta(hours=4)
    title = utc_now.strftime("%B %d, %Y") + "  |  " + utc_now.strftime("%H:%M UTC") + "  |  " + dc_time.strftime("%I:%M %p") + " Washington DC"

    log.info(f"=== SONNET REPORT START | {title} ===")

    try:
        sb     = get_supabase()
        client = get_anthropic()
    except Exception as e:
        return {"status": "ERROR", "error": str(e)}

    s2  = fetch_s2_full(sb)
    s3a = fetch_s3a(sb)
    s3d = fetch_s3d(sb)
    ma  = fetch_ma(sb)

    log.info(f"Data: S2={len(s2)} | S3-A={'yes' if s3a else 'no'} | S3-D={'yes' if s3d else 'no'} | MA={'yes' if ma else 'no'}")

    prompt = build_prompt(s2, s3a, s3d, ma)
    report = call_sonnet(client, prompt)
    if not report:
        return {"status": "GENERATION_FAILED"}

    docx_path = os.path.join(tempfile.gettempdir(),
        f"lens_{utc_now.strftime('%Y%m%d_%H%M')}.docx")

    docx_ok = generate_docx(report, docx_path, title)

    if docx_ok:
        send_telegram(docx_path, title)
    else:
        import requests
        token   = os.environ.get("TELEGRAM_BOT_TOKEN","")
        chat_id = os.environ.get("TELEGRAM_CHAT_ID","")
        if token and chat_id:
            chunks = [report[i:i+4000] for i in range(0, min(len(report),20000), 4000)]
            for chunk in chunks:
                requests.post(f"https://api.telegram.org/bot{token}/sendMessage",
                    json={"chat_id": chat_id, "text": chunk}, timeout=10)
                time.sleep(1)

    elapsed = round(time.time() - start, 1)
    log.info(f"=== DONE | {len(report.split())} words | {elapsed}s ===")
    return {"status": "OK", "words": len(report.split()), "docx": docx_ok}


if __name__ == "__main__":
    from dotenv import load_dotenv
    load_dotenv()
    result = run_sonnet_report()
    print(json.dumps(result, indent=2))# ── DOCX generation (pure python-docx) ──────────────────────────────────────────

def generate_docx(report_text: str, output_path: str, title: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        sec = doc.sections[0]
        sec.page_height   = Cm(29.7)
        sec.page_width    = Cm(21.0)
        sec.top_margin    = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin   = Cm(2.54)
        sec.right_margin  = Cm(2.54)

        # Page number in footer
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        for tag, ftype in [('w:fldChar','begin'), ('w:instrText', None), ('w:fldChar','end')]:
            el = OxmlElement(tag)
            if ftype:   el.set(qn('w:fldCharType'), ftype)
            else:       el.text = 'PAGE'
            run._r.append(el)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88,0x88,0x88)

        # Title
        tp = doc.add_paragraph()
        tr = tp.add_run(title)
        tr.bold = True
        tr.font.size = Pt(16)
        tr.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
        pPr = tp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'4');    bot.set(qn('w:color'),'1a1a2e')
        pBdr.append(bot); pPr.append(pBdr)
        doc.add_paragraph()

        for line in report_text.split('\n'):
            s = line.strip()
            if not s:
                doc.add_paragraph(); continue

            is_section = (s.endswith(':') and len(s.split()) >= 2 and
                s.rstrip(':').replace(' ','').replace('—','').replace('-','').isupper()
                and len(s) < 80)
            is_bullet  = s.startswith('- ')
            is_sub     = (s.endswith(':') and len(s) < 70 and
                not is_section and s[0].isupper())

            if is_section:
                p = doc.add_paragraph()
                r = p.add_run(s.rstrip(':'))
                r.bold = True; r.font.size = Pt(12)
                r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
                pP = p._p.get_or_add_pPr()
                pB = OxmlElement('w:pBdr')
                b2 = OxmlElement('w:bottom')
                b2.set(qn('w:val'),'single'); b2.set(qn('w:sz'),'2')
                b2.set(qn('w:space'),'4');    b2.set(qn('w:color'),'cccccc')
                pB.append(b2); pP.append(pB)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after  = Pt(6)
            elif is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(s[2:]).font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
            elif is_sub:
                p = doc.add_paragraph()
                r = p.add_run(s); r.bold = True; r.font.size = Pt(10)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(2)
            else:
                p = doc.add_paragraph(s)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(4)
                for r in p.runs: r.font.size = Pt(10)

        doc.save(output_path)
        log.info(f"DOCX generated: {output_path}")
        return True

    except Exception as e:
        log.error(f"DOCX failed: {e}")
        return False


# ── Telegram send ─────────────────────────────────────────────────────────────

def send_telegram(docx_path: str, title: str) -> bool:
    import requests
    token   = os.environ.get("TELEGRAM_BOT_TOKEN","")
    chat_id = os.environ.get("TELEGRAM_CHAT_ID","")
    if not token or not chat_id: return False
    try:
        fname = title.replace(' ','_').replace(',','').replace(':','') + ".docx"
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        with open(docx_path,'rb') as f:
            resp = requests.post(url,
                data={"chat_id": chat_id},
                files={"document": (fname, f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=60)
        ok = resp.status_code == 200
        log.info(f"Telegram: {'OK' if ok else 'FAILED'}")
        return ok
    except Exception as e:
        log.error(f"Telegram failed: {e}")
        return False


# ── Main ──────────────────────────────────────────────────────────────────────

def run_sonnet_report() -> dict:
    start = time.time()

    # Thai time for title (UTC+7)
    utc_now = datetime.now(timezone.utc)
    dc_time = utc_now - timedelta(hours=4)
    title = utc_now.strftime("%B %d, %Y") + "  |  " + utc_now.strftime("%H:%M UTC") + "  |  " + dc_time.strftime("%I:%M %p") + " Washington DC"

    log.info(f"=== SONNET REPORT START | {title} ===")

    try:
        sb     = get_supabase()
        client = get_anthropic()
    except Exception as e:
        return {"status": "ERROR", "error": str(e)}

    s2  = fetch_s2_full(sb)
    s3a = fetch_s3a(sb)
    s3d = fetch_s3d(sb)
    ma  = fetch_ma(sb)

    log.info(f"Data: S2={len(s2)} | S3-A={'yes' if s3a else 'no'} | S3-D={'yes' if s3d else 'no'} | MA={'yes' if ma else 'no'}")

    prompt = build_prompt(s2, s3a, s3d, ma)
    report = call_sonnet(client, prompt)
    if not report:
        return {"status": "GENERATION_FAILED"}

    docx_path = os.path.join(tempfile.gettempdir(),
        f"lens_{utc_now.strftime('%Y%m%d_%H%M')}.docx")

    docx_ok = generate_docx(report, docx_path, title)

    if docx_ok:
        send_telegram(docx_path, title)
    else:
        import requests
        token   = os.environ.get("TELEGRAM_BOT_TOKEN","")
        chat_id = os.environ.get("TELEGRAM_CHAT_ID","")
        if token and chat_id:
            chunks = [report[i:i+4000] for i in range(0, min(len(report),20000), 4000)]
            for chunk in chunks:
                requests.post(f"https://api.telegram.org/bot{token}/sendMessage",
                    json={"chat_id": chat_id, "text": chunk}, timeout=10)
                time.sleep(1)

    elapsed = round(time.time() - start, 1)
    log.info(f"=== DONE | {len(report.split())} words | {elapsed}s ===")
    return {"status": "OK", "words": len(report.split()), "docx": docx_ok}


if __name__ == "__main__":
    from dotenv import load_dotenv
    load_dotenv()
    result = run_sonnet_report()
    print(json.dumps(result, indent=2))
